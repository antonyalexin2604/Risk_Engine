#!/usr/bin/env python3
"""
Convert all .md files in Requirements folder to professional .docx format
with proper formatting, tables, code blocks, and styling.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level):
    """Add heading with proper styling based on level"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code_text, language=""):
    """Add formatted code block"""
    p = doc.add_paragraph(code_text, style='Normal')
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
    # Add gray background
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E0E0E0')
    p._element.get_or_add_pPr().append(shading_elm)
    return p

def add_table_from_markdown(doc, table_lines):
    """Convert markdown table to Word table"""
    rows = []
    for line in table_lines:
        if not line.strip() or line.strip().startswith('|') == False:
            continue
        # Remove leading/trailing pipes and split
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells and not all(c == '' or c == '-' * len(c) or re.match(r'^-+$', c) for c in cells):
            rows.append(cells)

    if not rows:
        return None

    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Fill cells
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data

    return table

def parse_and_convert_markdown(md_filepath, docx_filepath):
    """Parse markdown file and convert to DOCX"""
    # Read markdown file
    with open(md_filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Add title page metadata
    doc.add_heading('Prometheus Requirements Documentation', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown line by line
    lines = content.split('\n')
    i = 0
    current_table_lines = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            code_block = []
            language = line.replace('```', '').strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_block.append(lines[i])
                i += 1
            add_code_block(doc, '\n'.join(code_block), language)
            i += 1
            current_table_lines = []
            continue

        # Handle headings
        if line.startswith('#'):
            # Add any pending table
            if current_table_lines:
                add_table_from_markdown(doc, current_table_lines)
                current_table_lines = []

            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            if level <= 6:
                add_heading_with_style(doc, title, min(level, 3))
            i += 1
            continue

        # Handle tables
        if line.strip().startswith('|'):
            current_table_lines.append(line)
            i += 1
            continue

        # Flush table if line doesn't continue it
        if current_table_lines and not line.strip().startswith('|'):
            add_table_from_markdown(doc, current_table_lines)
            current_table_lines = []

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip() and not line.strip().startswith('#') and not line.strip().startswith('|'):
            p = doc.add_paragraph(line.strip())
            # Apply formatting for bold/italic
            if line.strip():
                # Replace **text** with bold
                if '**' in line:
                    parts = re.split(r'\*\*(.*?)\*\*', line.strip())
                    p.clear()
                    for idx, part in enumerate(parts):
                        run = p.add_run(part)
                        if idx % 2 == 1:  # Bold parts
                            run.bold = True

        i += 1

    # Flush any remaining table
    if current_table_lines:
        add_table_from_markdown(doc, current_table_lines)

    # Add footer with metadata
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Generated from: {Path(md_filepath).name}\n").italic = True
    footer.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    footer.add_run("Prometheus Phase 2A Requirements Documentation").italic = True

    # Save document
    doc.save(docx_filepath)
    return docx_filepath

def convert_all_markdown_files(requirements_dir):
    """Convert all .md files in directory to .docx"""
    requirements_path = Path(requirements_dir)

    if not requirements_path.exists():
        print(f"❌ Requirements directory not found: {requirements_dir}")
        return False

    md_files = sorted(requirements_path.glob('*.md'))

    if not md_files:
        print(f"❌ No .md files found in {requirements_dir}")
        return False

    print(f"🔄 Found {len(md_files)} .md files to convert\n")
    print("=" * 80)

    successful = 0
    failed = 0

    for md_file in md_files:
        docx_file = md_file.with_suffix('.docx')

        try:
            print(f"⏳ Converting: {md_file.name}")
            print(f"   → {docx_file.name}")

            parse_and_convert_markdown(str(md_file), str(docx_file))

            file_size = docx_file.stat().st_size
            print(f"   ✅ Success ({file_size:,} bytes)\n")
            successful += 1

        except Exception as e:
            print(f"   ❌ FAILED: {str(e)}\n")
            failed += 1

    print("=" * 80)
    print(f"\n📊 Conversion Summary:")
    print(f"   ✅ Successful: {successful}/{len(md_files)}")
    print(f"   ❌ Failed: {failed}/{len(md_files)}")

    if successful == len(md_files):
        print(f"\n✨ All .md files successfully converted to .docx!")
        print(f"📂 Location: {requirements_dir}")
        return True
    else:
        print(f"\n⚠️  {failed} file(s) failed to convert")
        return False

if __name__ == '__main__':
    requirements_dir = '/Users/aaron/Documents/Project/Prometheus/docs/Requirements'
    convert_all_markdown_files(requirements_dir)


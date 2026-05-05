#!/usr/bin/env python3
"""
Improved: Convert all .md files in Requirements folder to professional .docx format
with better error handling for edge cases, complex tables, and large files.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level):
    """Add heading with proper styling based on level"""
    heading = doc.add_heading(text, level=min(level, 3))
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code_text, language=""):
    """Add formatted code block with proper styling"""
    # Split into multiple paragraphs if too long
    lines = code_text.split('\n')
    for line in lines:
        if line.strip():
            p = doc.add_paragraph(line, style='Normal')
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(64, 64, 64)
            # Add light gray background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            p._element.get_or_add_pPr().append(shading_elm)
    return None

def add_table_from_markdown(doc, table_lines):
    """Convert markdown table to Word table with robust error handling"""
    if not table_lines:
        return None

    rows = []

    for line in table_lines:
        if not line.strip():
            continue

        # Skip separator lines (all dashes and pipes)
        if re.match(r'^[\s|\-]+$', line):
            continue

        # Extract cells from line
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Remove leading/trailing pipes
            clean_line = line.strip()[1:-1]
            cells = [cell.strip() for cell in clean_line.split('|')]

            # Filter out empty rows and separator rows
            if cells and len(cells) > 0:
                # Verify not all cells are dashes
                if not all(re.match(r'^-+$', c) or c == '' for c in cells):
                    rows.append(cells)

    if not rows or len(rows) < 2:
        # If no valid table found, just add as text
        for line in table_lines:
            if line.strip() and not re.match(r'^[\s|\-]+$', line):
                doc.add_paragraph(line.strip())
        return None

    try:
        # Ensure all rows have same number of columns
        max_cols = max(len(row) for row in rows) if rows else 1
        normalized_rows = []
        for row in rows:
            if len(row) < max_cols:
                row = row + [''] * (max_cols - len(row))
            normalized_rows.append(row[:max_cols])

        # Create table
        table = doc.add_table(rows=len(normalized_rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'

        # Fill cells and format header row
        for row_idx, row_data in enumerate(normalized_rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_data

                # Format header row (first row)
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)

        return table

    except Exception as e:
        print(f"      ⚠️  Table parsing error (skipping table): {str(e)}")
        return None

def parse_and_convert_markdown(md_filepath, docx_filepath):
    """Parse markdown file and convert to DOCX with improved error handling"""

    try:
        # Read markdown file
        with open(md_filepath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Create document
        doc = Document()

        # Add metadata
        title = doc.add_heading('Prometheus Requirements Documentation', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse markdown line by line
        lines = content.split('\n')
        i = 0
        current_table_lines = []
        in_code_block = False
        code_language = ""
        code_lines = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                # Flush any pending table
                if current_table_lines:
                    add_table_from_markdown(doc, current_table_lines)
                    current_table_lines = []

                if not in_code_block:
                    # Starting code block
                    in_code_block = True
                    code_language = line.replace('```', '').strip()
                    code_lines = []
                else:
                    # Ending code block
                    in_code_block = False
                    if code_lines:
                        add_code_block(doc, '\n'.join(code_lines), code_language)
                    code_lines = []

                i += 1
                continue

            # Collect code block lines
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Handle headings
            if line.startswith('#') and not in_code_block:
                # Flush any pending table
                if current_table_lines:
                    add_table_from_markdown(doc, current_table_lines)
                    current_table_lines = []

                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                if level <= 6 and title:
                    add_heading_with_style(doc, title, level)

                i += 1
                continue

            # Handle horizontal rules
            if re.match(r'^[\s\-\*]{3,}$', line.strip()):
                if current_table_lines:
                    add_table_from_markdown(doc, current_table_lines)
                    current_table_lines = []
                i += 1
                continue

            # Handle tables
            if line.strip().startswith('|') and '|' in line:
                current_table_lines.append(line)
                i += 1
                continue

            # Flush table if line doesn't continue it
            if current_table_lines and not (line.strip().startswith('|') and '|' in line):
                add_table_from_markdown(doc, current_table_lines)
                current_table_lines = []

            # Handle bullet lists
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:].strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue

            # Handle numbered lists
            if re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                if text:
                    doc.add_paragraph(text, style='List Number')
                i += 1
                continue

            # Handle regular paragraphs
            if line.strip() and not line.strip().startswith('#'):
                text = line.strip()

                # Check for bold markers
                if '**' in text or '__' in text or '*' in text:
                    p = doc.add_paragraph()

                    # Simple bold/italic handling
                    parts = re.split(r'(\*\*.*?\*\*|__.*?__|_.*?_|\*.*?\*)', text)

                    for part in parts:
                        if not part:
                            continue

                        if part.startswith('**') and part.endswith('**'):
                            # Bold
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        elif part.startswith('__') and part.endswith('__'):
                            # Bold (alt)
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        elif part.startswith('_') and part.endswith('_'):
                            # Italic
                            run = p.add_run(part[1:-1])
                            run.italic = True
                        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
                            # Italic (alt)
                            run = p.add_run(part[1:-1])
                            run.italic = True
                        else:
                            p.add_run(part)
                else:
                    doc.add_paragraph(text)

            i += 1

        # Flush any remaining elements
        if current_table_lines:
            add_table_from_markdown(doc, current_table_lines)

        if in_code_block and code_lines:
            add_code_block(doc, '\n'.join(code_lines), code_language)

        # Add footer with metadata
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"Source: {Path(md_filepath).name}\n").italic = True
        footer.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
        footer.add_run("Prometheus Phase 2A Documentation").italic = True
        footer_format = footer.paragraph_format
        footer_format.left_indent = Inches(0.5)

        # Save document
        doc.save(docx_filepath)
        return True

    except Exception as e:
        print(f"      ❌ Error: {str(e)}")
        return False

def convert_all_markdown_files(requirements_dir):
    """Convert all .md files in directory to .docx"""
    requirements_path = Path(requirements_dir)

    if not requirements_path.exists():
        print(f"❌ Requirements directory not found: {requirements_dir}")
        return False

    md_files = sorted(requirements_path.glob('*.md'))

    if not md_files:
        print(f"❌ No .md files found in {requirements_dir}")
        return False

    print(f"🔄 Converting {len(md_files)} .md files to .docx format\n")
    print("=" * 80)

    successful = 0
    failed = 0
    failed_files = []

    for md_file in md_files:
        docx_file = md_file.with_suffix('.docx')

        print(f"⏳ Converting: {md_file.name}")
        print(f"   → {docx_file.name}")

        success = parse_and_convert_markdown(str(md_file), str(docx_file))

        if success:
            try:
                file_size = docx_file.stat().st_size
                print(f"   ✅ Success ({file_size:,} bytes)\n")
                successful += 1
            except:
                print(f"   ✅ Conversion complete\n")
                successful += 1
        else:
            print(f"   ❌ FAILED\n")
            failed += 1
            failed_files.append(md_file.name)

    print("=" * 80)
    print(f"\n📊 Conversion Summary:")
    print(f"   ✅ Successful: {successful}/{len(md_files)}")
    print(f"   ❌ Failed: {failed}/{len(md_files)}")

    if failed_files:
        print(f"\n   Failed files:")
        for fname in failed_files:
            print(f"      • {fname}")

    if successful == len(md_files):
        print(f"\n✨ All {len(md_files)} .md files successfully converted to .docx!")
        print(f"📂 Location: {requirements_dir}")
        return True
    else:
        print(f"\n⚠️  {failed} file(s) failed to convert")
        return failed == 0

if __name__ == '__main__':
    requirements_dir = '/Users/aaron/Documents/Project/Prometheus/docs/Requirements'
    success = convert_all_markdown_files(requirements_dir)
    exit(0 if success else 1)


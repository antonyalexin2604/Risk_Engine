#!/usr/bin/env python3
"""Convert AIRB_TECHNICAL_GUIDE.md to AIRB_TECHNICAL_GUIDE.docx"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC  = os.path.join(BASE, 'docs', 'AIRB_TECHNICAL_GUIDE.md')
DST  = os.path.join(BASE, 'docs', 'AIRB_TECHNICAL_GUIDE.docx')

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)


def add_heading(text, level):
    colors = {1: '1F3864', 2: '2E5797', 3: '2E75B6', 4: '4472C4'}
    sizes  = {1: 20,       2: 16,       3: 14,       4: 13}
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(colors.get(level, '000000'))
        run.font.size = Pt(sizes.get(level, 12))
        run.font.bold = True
    return para


def add_code(text):
    para = doc.add_paragraph()
    para.style = doc.styles['No Spacing']
    run = para.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  'EBF3FB')
    para._p.get_or_add_pPr().append(shading)
    return para


def inline_format(para, text):
    """Render inline `code` and **bold** into a paragraph."""
    for part in re.split(r'(`[^`]+`|\*\*[^*]+\*\*)', text):
        if part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name  = 'Courier New'
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)
        elif part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code = False
code_lines = []

while i < len(lines):
    raw     = lines[i].rstrip('\n')
    stripped = raw.strip()

    # ── code fence ────────────────────────────────────────────────
    if stripped.startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            add_code('\n'.join(code_lines))
        i += 1
        continue

    if in_code:
        code_lines.append(raw)
        i += 1
        continue

    # ── heading ───────────────────────────────────────────────────
    m = re.match(r'^(#{1,4})\s+(.*)', stripped)
    if m:
        add_heading(m.group(2), len(m.group(1)))
        i += 1
        continue

    # ── horizontal rule ───────────────────────────────────────────
    if stripped in ('---', '***', '___'):
        p = doc.add_paragraph('─' * 80)
        p.style = doc.styles['No Spacing']
        i += 1
        continue

    # ── table ─────────────────────────────────────────────────────
    if stripped.startswith('|'):
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
            if not all(re.match(r'^[-:]+$', c) for c in cells if c):
                rows.append(cells)
            i += 1
        if rows:
            max_cols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=max_cols)
            tbl.style = 'Light List Accent 1'
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < max_cols:
                        cell = tbl.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        inline_format(p, cell_text)
                        if ri == 0:
                            for run in p.runs:
                                run.bold = True
        doc.add_paragraph()
        continue

    # ── blockquote ────────────────────────────────────────────────
    if stripped.startswith('>'):
        content = stripped.lstrip('> ').strip()
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.4)
        run = para.add_run(content)
        run.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        i += 1
        continue

    # ── bullet list ───────────────────────────────────────────────
    if re.match(r'^[-*]\s', stripped):
        para = doc.add_paragraph(style='List Bullet')
        inline_format(para, stripped[2:])
        i += 1
        continue

    # ── numbered list ─────────────────────────────────────────────
    m = re.match(r'^\d+\.\s+(.*)', stripped)
    if m:
        para = doc.add_paragraph(style='List Number')
        inline_format(para, m.group(1))
        i += 1
        continue

    # ── blank line ────────────────────────────────────────────────
    if stripped == '':
        doc.add_paragraph()
        i += 1
        continue

    # ── normal paragraph ──────────────────────────────────────────
    para = doc.add_paragraph()
    inline_format(para, stripped)
    i += 1

doc.save(DST)
print(f"Saved: {DST}")


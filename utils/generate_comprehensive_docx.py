"""
PROMETHEUS — Generate Comprehensive Project Documentation (Word .docx)
Combines FSD, IMM Technical Guide, Market Data Architecture, CVA Enhancements,
FRTB FRD Summary, and README into a single institutional-grade Word document.

Output: PROMETHEUS_COMPREHENSIVE_DOCUMENTATION.docx
"""

from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Brand Colours ────────────────────────────────────────────────────────────
DEEP_SLATE   = RGBColor(0x1E, 0x2A, 0x3A)   # #1E2A3A — headings, table headers
CRIMSON      = RGBColor(0x9B, 0x11, 0x1E)   # #9B111E — accent
GOLD         = RGBColor(0xC8, 0x96, 0x20)   # #C89620 — part titles
LIGHT_SLATE  = RGBColor(0x4A, 0x5A, 0x6A)   # #4A5A6A — body text
MID_GREY     = RGBColor(0xD0, 0xD4, 0xD9)   # table borders
CODE_BG      = RGBColor(0xF4, 0xF4, 0xF4)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

BASE_DIR = Path("/Users/aaron/Documents/Project/Prometheus")
OUTPUT   = BASE_DIR / "docs" / "PROMETHEUS_COMPREHENSIVE_DOCUMENTATION.docx"

# Source files in document order (all under docs/)
SOURCES = [
    ("PART I — FUNCTIONAL SPECIFICATION DOCUMENT",
     BASE_DIR / "docs" / "PROMETHEUS_FSD.md"),
    ("PART II — IMM TECHNICAL REFERENCE GUIDE",
     BASE_DIR / "docs" / "IMM_TECHNICAL_GUIDE.md"),
    ("PART III — MARKET DATA ARCHITECTURE",
     BASE_DIR / "docs" / "MARKET_DATA_ARCHITECTURE.md"),
    ("PART IV — CVA ENHANCEMENTS & MAR50 COMPLIANCE",
     BASE_DIR / "docs" / "CVA_ENHANCEMENTS_SUMMARY.md"),
    ("PART V — CVA IMPLEMENTATION GUIDE",
     BASE_DIR / "docs" / "CVA_IMPLEMENTATION_GUIDE.md"),
]


# ─── XML / Style Helpers ──────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs) -> None:
    tc        = cell._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        cfg = kwargs.get(edge, {"sz": "4", "val": "single", "color": "C8CDD3"})
        tag.set(qn("w:val"),   cfg.get("val",   "single"))
        tag.set(qn("w:sz"),    cfg.get("sz",    "4"))
        tag.set(qn("w:color"), cfg.get("color", "C8CDD3"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc: Document, color_hex: str = "9B111E") -> None:
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)


def add_thin_rule(doc: Document, color_hex: str = "D0D4D9") -> None:
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def _parse_inline(p, text: str, base_size: float = 10.5,
                  base_color: RGBColor = LIGHT_SLATE) -> None:
    """Parse **bold**, *italic*, `code` markers into runs on paragraph p."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run            = p.add_run(part[2:-2])
            run.bold       = True
            run.font.size  = Pt(base_size)
            run.font.color.rgb = DEEP_SLATE
        elif part.startswith("`") and part.endswith("`"):
            run                = p.add_run(part[1:-1])
            run.font.name      = "Courier New"
            run.font.size      = Pt(base_size - 1.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith("*") and part.endswith("*"):
            run            = p.add_run(part[1:-1])
            run.italic     = True
            run.font.size  = Pt(base_size)
            run.font.color.rgb = base_color
        else:
            run                = p.add_run(part)
            run.font.size      = Pt(base_size)
            run.font.color.rgb = base_color


def add_body_paragraph(doc: Document, text: str, indent: int = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent == 1:
        p.paragraph_format.left_indent = Inches(0.3)
    elif indent == 2:
        p.paragraph_format.left_indent = Inches(0.6)
    _parse_inline(p, text)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    _parse_inline(p, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F4F4F4")
        pPr.append(shd)
        run                = p.add_run(line if line else " ")
        run.font.name      = "Courier New"
        run.font.size      = Pt(8.5)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_table_from_md(doc: Document, md_rows: list[str]) -> None:
    data_rows = [r for r in md_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    n_cols = max(len(r) for r in parsed)
    n_rows = len(parsed)

    tbl           = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    usable_width = Inches(6.0)
    col_w        = usable_width / n_cols
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_w

    for i, row_data in enumerate(parsed):
        row       = tbl.rows[i]
        row.height = Cm(0.65)
        is_header = (i == 0)
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell                    = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if is_header:
                set_cell_bg(cell, "1E2A3A")
            elif i % 2 == 0:
                set_cell_bg(cell, "F0F4F8")
            else:
                set_cell_bg(cell, "FFFFFF")

            for k in ("top", "left", "bottom", "right"):
                set_cell_border(cell, **{k: {"val": "single", "sz": "4", "color": "C8CDD3"}})

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Pt(4)

            clean   = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', cell_text)
            pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
            parts   = pattern.split(clean)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run                = p.add_run(part[2:-2])
                    run.bold           = True
                    run.font.size      = Pt(9.5)
                    run.font.color.rgb = WHITE if is_header else DEEP_SLATE
                elif part.startswith("`") and part.endswith("`"):
                    run                = p.add_run(part[1:-1])
                    run.font.name      = "Courier New"
                    run.font.size      = Pt(8.5)
                    run.font.color.rgb = WHITE if is_header else RGBColor(0xC7, 0x25, 0x4E)
                elif part.startswith("*") and part.endswith("*"):
                    run                = p.add_run(part[1:-1])
                    run.italic         = True
                    run.font.size      = Pt(9.5)
                    run.font.color.rgb = WHITE if is_header else LIGHT_SLATE
                else:
                    run                = p.add_run(part)
                    run.font.size      = Pt(9.5)
                    run.font.color.rgb = WHITE if is_header else LIGHT_SLATE
                if is_header:
                    run.bold = True

    doc.add_paragraph()


# ─── Document Setup ───────────────────────────────────────────────────────────

def setup_document() -> Document:
    doc = Document()

    for section in doc.sections:
        section.page_width    = Inches(8.27)
        section.page_height   = Inches(11.69)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name      = "Calibri"
    normal.font.size      = Pt(10.5)
    normal.font.color.rgb = LIGHT_SLATE
    normal.paragraph_format.space_after = Pt(6)

    h1 = styles["Heading 1"]
    h1.font.name                       = "Calibri"
    h1.font.size                       = Pt(20)
    h1.font.bold                       = True
    h1.font.color.rgb                  = DEEP_SLATE
    h1.paragraph_format.space_before   = Pt(18)
    h1.paragraph_format.space_after    = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name                       = "Calibri"
    h2.font.size                       = Pt(15)
    h2.font.bold                       = True
    h2.font.color.rgb                  = CRIMSON
    h2.paragraph_format.space_before   = Pt(14)
    h2.paragraph_format.space_after    = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name                       = "Calibri"
    h3.font.size                       = Pt(12)
    h3.font.bold                       = True
    h3.font.color.rgb                  = DEEP_SLATE
    h3.paragraph_format.space_before   = Pt(10)
    h3.paragraph_format.space_after    = Pt(3)
    h3.paragraph_format.keep_with_next = True

    h4 = styles["Heading 4"]
    h4.font.name                       = "Calibri"
    h4.font.size                       = Pt(11)
    h4.font.bold                       = True
    h4.font.italic                     = True
    h4.font.color.rgb                  = LIGHT_SLATE
    h4.paragraph_format.space_before   = Pt(8)
    h4.paragraph_format.space_after    = Pt(2)

    return doc


# ─── Cover Page ───────────────────────────────────────────────────────────────

def add_cover_page(doc: Document) -> None:
    for _ in range(5):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    # Trident symbol line
    sym = doc.add_paragraph()
    sym.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr  = sym.add_run("🔱")
    sr.font.size  = Pt(28)

    # Main title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PROMETHEUS")
    r.font.name      = "Calibri"
    r.font.size      = Pt(42)
    r.font.bold      = True
    r.font.color.rgb = DEEP_SLATE

    # Subtitle
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("BASEL III/IV RISK MANAGEMENT PLATFORM")
    r2.font.name      = "Calibri"
    r2.font.size      = Pt(18)
    r2.font.color.rgb = CRIMSON
    r2.font.bold      = True

    doc.add_paragraph()
    add_horizontal_rule(doc, "9B111E")
    doc.add_paragraph()

    # Document type
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = dt.add_run("Comprehensive Project Documentation")
    r3.font.name      = "Calibri"
    r3.font.size      = Pt(17)
    r3.font.color.rgb = DEEP_SLATE

    # Secondary subtitle
    ds = doc.add_paragraph()
    ds.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = ds.add_run(
        "Functional Specification · Technical Reference · Architecture · Regulatory Compliance"
    )
    r4.font.name      = "Calibri"
    r4.font.size      = Pt(12)
    r4.font.italic    = True
    r4.font.color.rgb = LIGHT_SLATE

    doc.add_paragraph()

    # Meta table
    meta = [
        ("Document Reference",  "PROMETHEUS-CPD-v1.0"),
        ("Document Type",       "Comprehensive Project Documentation"),
        ("Version",             "1.0"),
        ("Date",                "April 14, 2026"),
        ("Classification",      "Internal Use — Confidential"),
        ("Status",              "Final"),
        ("Platform",            "Python 3.11 · PostgreSQL 15 · Streamlit 1.28"),
        ("Prepared By",         "Risk Technology — Lead Developer"),
        ("Reviewed By",         "Head of Market Risk | Head of Credit Risk | Regulatory Affairs"),
        ("Regulatory Basis",    "Basel III/IV — BCBS Standards CRE52, CRE53, CRE30–36, MAR20–50, RBC20, OPE25"),
    ]

    tbl           = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        rc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        rc.width = Inches(2.5)
        vc.width = Inches(3.8)
        set_cell_bg(rc, "1E2A3A")
        set_cell_bg(vc, "F0F4F8" if i % 2 == 0 else "FFFFFF")
        set_cell_border(rc)
        set_cell_border(vc)
        rr = rc.paragraphs[0].add_run(k)
        rr.font.name      = "Calibri"
        rr.font.size      = Pt(10)
        rr.bold           = True
        rr.font.color.rgb = WHITE
        vr = vc.paragraphs[0].add_run(v)
        vr.font.name      = "Calibri"
        vr.font.size      = Pt(10)
        vr.font.color.rgb = DEEP_SLATE

    doc.add_paragraph()
    add_thin_rule(doc)

    # Scope summary
    scope_p = doc.add_paragraph()
    scope_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = scope_p.add_run(
        "This document provides the complete project specification, technical design, "
        "regulatory traceability, and implementation guide for the PROMETHEUS platform — "
        "covering SA-CCR, IMM, A-IRB, FRTB, CVA, CCP, and Operational Risk engines."
    )
    sr.font.name      = "Calibri"
    sr.font.size      = Pt(9.5)
    sr.font.italic    = True
    sr.font.color.rgb = LIGHT_SLATE

    doc.add_page_break()


# ─── Document Outline Page ────────────────────────────────────────────────────

def add_outline_page(doc: Document) -> None:
    p = doc.add_paragraph("DOCUMENT OUTLINE", style="Heading 1")
    add_horizontal_rule(doc)

    parts = [
        ("PART I",   "Functional Specification Document (FSD)",
         "Platform overview, project objectives, functional requirements for all six capital engines, "
         "business workflow, system architecture, technical design, database schema, and appendices."),
        ("PART II",  "IMM Technical Reference Guide",
         "Monte Carlo simulation framework, EPE/EEPE exposure metrics, stochastic process models "
         "(GBM + Hull-White), stressed EEPE calibration, CSA collateral adjustment, CVA linkage, "
         "sensitivity analysis, and complete glossary."),
        ("PART III", "Market Data Architecture",
         "Pluggable market data provider interface, Bloomberg/Refinitiv/Internal/Static implementations, "
         "CDSSpreadService with fallback and caching, configuration, usage examples, and production checklist."),
        ("PART IV",  "CVA Enhancements & MAR50 Compliance",
         "SA-CVA six risk-class framework, vega charge implementation, proxy spread monthly review workflow, "
         "capital floor verification, and compliance summary against MAR50.43/44/48/49 and CAP10 FAQ1."),
        ("PART V",   "CVA Implementation Guide",
         "Current implementation status, remaining integration tasks, production deployment checklist, "
         "and regulatory sign-off preparation steps for SA-CVA supervisory approval."),
    ]

    tbl           = doc.add_table(rows=len(parts) + 1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    headers = ["Part", "Title", "Contents"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, "1E2A3A")
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(h)
        r.font.name      = "Calibri"
        r.font.bold      = True
        r.font.size      = Pt(10.5)
        r.font.color.rgb = WHITE

    widths = [Inches(0.8), Inches(2.2), Inches(3.2)]
    for j, col in enumerate(tbl.columns):
        for cell in col.cells:
            cell.width = widths[j]

    for i, (part, title, contents) in enumerate(parts):
        row = tbl.rows[i + 1]
        row.height = Cm(1.8)
        data = [part, title, contents]
        bg   = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        for j, txt in enumerate(data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Pt(4)
            r = p.add_run(txt)
            r.font.name      = "Calibri"
            r.font.size      = Pt(9.5)
            r.font.color.rgb = DEEP_SLATE if j < 2 else LIGHT_SLATE
            if j == 0:
                r.bold = True

    doc.add_paragraph()

    add_body_paragraph(doc,
        "**Total coverage:** Six regulatory capital engines · 48/48 tests passing · "
        "8-page interactive dashboard · Full audit trail · ~2-minute daily risk run")

    doc.add_page_break()


# ─── Part Separator Page ──────────────────────────────────────────────────────

def add_part_separator(doc: Document, part_label: str, part_title: str,
                        description: str) -> None:
    doc.add_page_break()

    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    add_horizontal_rule(doc, "C89620")

    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pl.add_run(part_label)
    r1.font.name      = "Calibri"
    r1.font.size      = Pt(14)
    r1.font.bold      = True
    r1.font.color.rgb = GOLD

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = pt.add_run(part_title)
    r2.font.name      = "Calibri"
    r2.font.size      = Pt(24)
    r2.font.bold      = True
    r2.font.color.rgb = DEEP_SLATE

    doc.add_paragraph()
    add_horizontal_rule(doc, "C89620")
    doc.add_paragraph()

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = desc_p.add_run(description)
    r3.font.name      = "Calibri"
    r3.font.size      = Pt(10.5)
    r3.font.italic    = True
    r3.font.color.rgb = LIGHT_SLATE

    doc.add_page_break()


# ─── Markdown → DOCX Parser ───────────────────────────────────────────────────

def render_md_to_docx(doc: Document, md_path: Path,
                      skip_lines: int = 0) -> None:
    """
    Parse a Markdown file and render it into the Word document.
    skip_lines: number of header/metadata lines to skip at the start.
    """
    lines = md_path.read_text(encoding="utf-8").splitlines()

    in_code_block  = False
    code_lines: list[str] = []
    in_table       = False
    table_rows: list[str] = []
    line_idx       = 0

    while line_idx < len(lines):
        line = lines[line_idx]

        # Skip leading metadata lines
        if line_idx < skip_lines:
            line_idx += 1
            continue

        stripped = line.strip()

        # ── Code fence ──────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                add_code_block(doc, code_lines)
                code_lines = []
            else:
                if in_table:
                    add_table_from_md(doc, table_rows)
                    table_rows = []
                    in_table   = False
                in_code_block = True
            line_idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            line_idx += 1
            continue

        # ── Table rows ──────────────────────────────────────────────────────
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_rows.append(stripped)
            line_idx += 1
            continue
        elif in_table:
            add_table_from_md(doc, table_rows)
            table_rows = []
            in_table   = False
            # re-process current line without advancing

        # ── Horizontal rules ────────────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            add_horizontal_rule(doc)
            line_idx += 1
            continue

        # ── Headings ────────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if h_match:
            level  = len(h_match.group(1))
            h_text = h_match.group(2).strip()
            h_text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', h_text)
            h_text = h_text.replace("`", "")
            style  = {1: "Heading 1", 2: "Heading 2",
                      3: "Heading 3", 4: "Heading 4"}.get(level, "Heading 4")
            doc.add_paragraph(h_text, style=style)
            if level == 1:
                add_horizontal_rule(doc)
            line_idx += 1
            continue

        # ── Bullet lists ────────────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text   = bullet_match.group(2).strip()
            text   = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
            add_bullet(doc, text, level=indent)
            line_idx += 1
            continue

        # ── Numbered lists ──────────────────────────────────────────────────
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_match:
            indent = len(num_match.group(1)) // 2
            text   = num_match.group(2).strip()
            text   = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
            p      = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.25)
            _parse_inline(p, text)
            line_idx += 1
            continue

        # ── Multi-line LaTeX block $$ ... $$ ────────────────────────────────
        if stripped.startswith("$$"):
            formula_lines: list[str] = []
            if stripped != "$$":
                formula_lines.append(stripped)
                line_idx += 1
            else:
                line_idx += 1
                while line_idx < len(lines) and lines[line_idx].strip() != "$$":
                    formula_lines.append(lines[line_idx])
                    line_idx += 1
                line_idx += 1  # skip closing $$
            add_code_block(doc, formula_lines)
            line_idx += 1
            continue

        # ── Inline formula $...$ ────────────────────────────────────────────
        if stripped.startswith("$") and not stripped.startswith("$$"):
            add_body_paragraph(doc, stripped, indent=1)
            line_idx += 1
            continue

        # ── Empty lines ─────────────────────────────────────────────────────
        if not stripped:
            line_idx += 1
            continue

        # ── Body text ───────────────────────────────────────────────────────
        clean = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', stripped)
        add_body_paragraph(doc, clean)
        line_idx += 1

    # Flush
    if in_table:
        add_table_from_md(doc, table_rows)
    if in_code_block:
        add_code_block(doc, code_lines)


# ─── Header / Footer ──────────────────────────────────────────────────────────

def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True

        hdr = section.header
        hp  = hdr.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = hp.add_run(
            "PROMETHEUS  |  Comprehensive Project Documentation  |  CONFIDENTIAL"
        )
        r1.font.name      = "Calibri"
        r1.font.size      = Pt(8)
        r1.font.color.rgb = LIGHT_SLATE

        pPr  = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "D0D4D9")
        pBdr.append(bot)
        pPr.append(pBdr)

        ftr = section.footer
        fp  = ftr.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = fp.add_run(
            "PROMETHEUS-CPD-v1.0  ·  Internal Use — Confidential  ·  April 14, 2026"
        )
        r2.font.size      = Pt(8)
        r2.font.color.rgb = LIGHT_SLATE


# ─── Part Descriptions ────────────────────────────────────────────────────────

PART_DESCRIPTIONS = {
    "PART I — FUNCTIONAL SPECIFICATION DOCUMENT": (
        "PART I", "Functional Specification Document",
        "Project description, objectives, components, upstream/downstream dependencies, "
        "eleven functional requirements (SA-CCR · IMM · A-IRB · FRTB · CVA · CCP · "
        "Orchestrator · Market Data · Backtesting · Dashboard · Audit Trail), "
        "system architecture, technical design, and regulatory appendices."
    ),
    "PART II — IMM TECHNICAL REFERENCE GUIDE": (
        "PART II", "IMM Technical Reference Guide",
        "Complete technical reference for the Internal Models Method Monte Carlo engine — "
        "simulation framework, EPE/EEPE/EEE definitions, GBM & Hull-White stochastic models, "
        "stressed EEPE calibration (2007–09), CSA adjustment, CVA linkage, Greeks, "
        "and comprehensive glossary of all exposure metrics."
    ),
    "PART III — MARKET DATA ARCHITECTURE": (
        "PART III", "Market Data Architecture",
        "Pluggable market data provider architecture — Bloomberg Terminal, Refinitiv/LSEG Eikon, "
        "internal REST API, and static test providers. CDSSpreadService with automatic fallback "
        "chains and 5-minute in-memory caching. Configuration, integration examples, "
        "and production deployment checklist."
    ),
    "PART IV — CVA ENHANCEMENTS & MAR50 COMPLIANCE": (
        "PART IV", "CVA Enhancements & MAR50 Compliance",
        "Four critical CVA enhancements: SA-CVA six risk-class delta framework (MAR50.43), "
        "vega charge implementation (MAR50.48), proxy spread monthly review workflow (MAR50.32(3)), "
        "and capital floor exclusion verification (CAP10 FAQ1). Capital impact analysis and "
        "production deployment roadmap."
    ),
    "PART V — CVA IMPLEMENTATION GUIDE": (
        "PART V", "CVA Implementation Guide",
        "Current implementation status of SA-CVA risk classes, integration tasks for full "
        "MAR50 compliance, production deployment phases (proxy spread registry, vega charge, "
        "full sensitivity integration, regulatory sign-off), and testing requirements."
    ),
}


# ─── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    print("=" * 60)
    print("  PROMETHEUS — Comprehensive Documentation Generator")
    print("=" * 60)

    doc = setup_document()

    # Cover page
    print("\n  [1/7] Building cover page...")
    add_cover_page(doc)

    # Document outline
    print("  [2/7] Building document outline...")
    add_outline_page(doc)

    # Process each part
    step = 3
    for part_label, md_path in SOURCES:
        if not md_path.exists():
            print(f"  [WARN] Skipping missing file: {md_path}")
            continue

        desc_tuple = PART_DESCRIPTIONS.get(part_label)
        if desc_tuple:
            label, title, description = desc_tuple
        else:
            label, title, description = part_label, part_label, ""

        print(f"  [{step}/7] Rendering {label}: {title[:50]}...")

        add_part_separator(doc, label, title, description)

        # Skip front-matter lines: FSD has ~12 metadata lines; others ~6
        skip = 12 if "FSD" in md_path.name else 6
        render_md_to_docx(doc, md_path, skip_lines=skip)

        step += 1

    # Header / footer on all sections
    print("  [7/7] Adding headers and footers...")
    add_header_footer(doc)

    # Save
    doc.save(str(OUTPUT))
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"\n{'=' * 60}")
    print(f"  ✅  Document saved: {OUTPUT}")
    print(f"  📄  File size:      {size_kb:.1f} KB")
    print(f"  📚  Parts:          {len(SOURCES)}")
    print(f"{'=' * 60}\n")


if __name__ == "__main__":
    main()


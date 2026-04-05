"""
PROMETHEUS — Generate FSD Word Document
Converts PROMETHEUS_FSD.md into a fully formatted PROMETHEUS_FSD.docx
"""

from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Brand Colours ────────────────────────────────────────────────────────────
DEEP_SLATE   = RGBColor(0x1E, 0x2A, 0x3A)   # #1E2A3A — headings, table headers
CRIMSON      = RGBColor(0x9B, 0x11, 0x1E)   # #9B111E — accent, rule lines
WARM_WHITE   = RGBColor(0xFF, 0xFB, 0xF5)   # #FFFBF5 — page background hint
LIGHT_SLATE  = RGBColor(0x4A, 0x5A, 0x6A)   # #4A5A6A — body text
MID_GREY     = RGBColor(0xD0, 0xD4, 0xD9)   # #D0D4D9 — table borders
PALE_BLUE    = RGBColor(0xE8, 0xF0, 0xF8)   # table header fill
CODE_BG      = RGBColor(0xF4, 0xF4, 0xF4)   # code block background

SOURCE_MD    = Path("/Users/aaron/Documents/Project/Prometheus/PROMETHEUS_FSD.md")
OUTPUT_DOCX  = Path("/Users/aaron/Documents/Project/Prometheus/PROMETHEUS_FSD.docx")


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    """Set table cell background colour via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set individual borders on a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag  = OxmlElement(f"w:{edge}")
        cfg  = kwargs.get(edge, {"sz": "4", "val": "single", "color": "D0D4D9"})
        tag.set(qn("w:val"),   cfg.get("val",   "single"))
        tag.set(qn("w:sz"),    cfg.get("sz",    "4"))
        tag.set(qn("w:color"), cfg.get("color", "D0D4D9"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc: Document, color_hex: str = "9B111E"):
    """Insert a thin horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p


def apply_run_inline(run, text: str):
    """Parse **bold** and `code` inline markers within a run's parent paragraph."""
    # This is applied at paragraph level in add_body_paragraph
    pass


def add_body_paragraph(doc: Document, text: str, indent: int = 0) -> None:
    """
    Add a body paragraph, parsing inline **bold**, *italic*, and `code` markers.
    indent: left indent level (0=none, 1=first, 2=second)
    """
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent == 1:
        p.paragraph_format.left_indent = Inches(0.3)
    elif indent == 2:
        p.paragraph_format.left_indent = Inches(0.6)

    # Parse inline markers
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run      = p.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = DEEP_SLATE
        elif part.startswith("`") and part.endswith("`"):
            run                   = p.add_run(part[1:-1])
            run.font.name         = "Courier New"
            run.font.size         = Pt(9)
            run.font.color.rgb    = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith("*") and part.endswith("*"):
            run        = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.size      = Pt(10.5)
        run.font.color.rgb = LIGHT_SLATE


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """Add a bullet list item with inline formatting."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)

    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run      = p.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = DEEP_SLATE
        elif part.startswith("`") and part.endswith("`"):
            run                = p.add_run(part[1:-1])
            run.font.name      = "Courier New"
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run                = p.add_run(part)
            run.font.size      = Pt(10.5)
            run.font.color.rgb = LIGHT_SLATE


def add_code_block(doc: Document, lines: list[str]) -> None:
    """Add a shaded monospace code block."""
    for line in lines:
        p                          = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        # shade via paragraph border/shading (approximated via XML)
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F4F4F4")
        pPr.append(shd)
        run                = p.add_run(line if line else " ")
        run.font.name      = "Courier New"
        run.font.size      = Pt(8.5)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def add_table_from_md(doc: Document, md_rows: list[str]) -> None:
    """
    Render a Markdown table (pipe-delimited rows) as a styled Word table.
    md_rows: list of raw '| col | col |' strings (includes separator row).
    """
    # Filter out separator row (e.g. |---|---|)
    data_rows = [r for r in md_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    n_cols = max(len(r) for r in parsed)
    n_rows = len(parsed)

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    # Column widths — distribute evenly across usable page width (~5.8 inches for A4 portrait)
    usable_width = Inches(6.0)
    col_w        = usable_width / n_cols
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_w

    for i, row_data in enumerate(parsed):
        row = tbl.rows[i]
        row.height = Cm(0.65)
        is_header  = (i == 0)
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Background
            if is_header:
                set_cell_bg(cell, "1E2A3A")
            elif i % 2 == 0:
                set_cell_bg(cell, "F0F4F8")
            else:
                set_cell_bg(cell, "FFFFFF")

            # Cell borders
            set_cell_border(cell, **{
                k: {"val": "single", "sz": "4", "color": "C8CDD3"}
                for k in ("top", "left", "bottom", "right")
            })

            # Cell content with inline formatting
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Pt(4)

            # Strip markdown links [text](url) → text
            clean = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', cell_text)
            # Parse inline bold/code
            pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
            parts   = pattern.split(clean)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run            = p.add_run(part[2:-2])
                    run.bold       = True
                    run.font.size  = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if is_header else DEEP_SLATE
                elif part.startswith("`") and part.endswith("`"):
                    run                = p.add_run(part[1:-1])
                    run.font.name      = "Courier New"
                    run.font.size      = Pt(8.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if is_header else RGBColor(0xC7, 0x25, 0x4E)
                else:
                    run            = p.add_run(part)
                    run.font.size  = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if is_header else LIGHT_SLATE
                if is_header:
                    run.bold = True

    doc.add_paragraph()  # spacing after table


# ─── Document Setup ───────────────────────────────────────────────────────────

def setup_document() -> Document:
    doc = Document()

    # Page margins (A4)
    for section in doc.sections:
        section.page_width      = Inches(8.27)
        section.page_height     = Inches(11.69)
        section.left_margin     = Inches(1.1)
        section.right_margin    = Inches(1.1)
        section.top_margin      = Inches(1.0)
        section.bottom_margin   = Inches(1.0)

    styles = doc.styles

    # ── Normal ──────────────────────────────────────────────────────────────
    normal = styles["Normal"]
    normal.font.name      = "Calibri"
    normal.font.size      = Pt(10.5)
    normal.font.color.rgb = LIGHT_SLATE
    normal.paragraph_format.space_after = Pt(6)

    # ── Heading 1 ──────────────────────────────────────────────────────────
    h1 = styles["Heading 1"]
    h1.font.name      = "Calibri"
    h1.font.size      = Pt(20)
    h1.font.bold      = True
    h1.font.color.rgb = DEEP_SLATE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # ── Heading 2 ──────────────────────────────────────────────────────────
    h2 = styles["Heading 2"]
    h2.font.name      = "Calibri"
    h2.font.size      = Pt(15)
    h2.font.bold      = True
    h2.font.color.rgb = CRIMSON
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # ── Heading 3 ──────────────────────────────────────────────────────────
    h3 = styles["Heading 3"]
    h3.font.name      = "Calibri"
    h3.font.size      = Pt(12)
    h3.font.bold      = True
    h3.font.color.rgb = DEEP_SLATE
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(3)
    h3.paragraph_format.keep_with_next = True

    # ── Heading 4 ──────────────────────────────────────────────────────────
    h4 = styles["Heading 4"]
    h4.font.name      = "Calibri"
    h4.font.size      = Pt(11)
    h4.font.bold      = True
    h4.font.italic    = True
    h4.font.color.rgb = LIGHT_SLATE
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after  = Pt(2)

    return doc


# ─── Cover Page ───────────────────────────────────────────────────────────────

def add_cover_page(doc: Document) -> None:
    # Spacer
    for _ in range(4):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    # Main title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PROMETHEUS")
    r.font.name      = "Calibri"
    r.font.size      = Pt(36)
    r.font.bold      = True
    r.font.color.rgb = DEEP_SLATE

    # Subtitle
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("RISK MANAGEMENT PLATFORM")
    r2.font.name      = "Calibri"
    r2.font.size      = Pt(18)
    r2.font.color.rgb = CRIMSON
    r2.font.bold      = True

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Document type
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = dt.add_run("Functional Specification Document")
    r3.font.name      = "Calibri"
    r3.font.size      = Pt(16)
    r3.font.color.rgb = DEEP_SLATE

    doc.add_paragraph()

    # Meta table
    meta = [
        ("Document Reference",  "PROMETHEUS-FSD-v1.0"),
        ("Version",             "1.0"),
        ("Date",                "April 5, 2026"),
        ("Classification",      "Internal Use — Confidential"),
        ("Status",              "Final"),
        ("Prepared By",         "Risk Technology — Lead Developer"),
        ("Reviewed By",         "Head of Market Risk | Head of Credit Risk | Regulatory Affairs"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        rc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        rc.width = Inches(2.2)
        vc.width = Inches(3.8)
        set_cell_bg(rc, "1E2A3A")
        set_cell_bg(vc, "F0F4F8" if i % 2 == 0 else "FFFFFF")
        set_cell_border(rc)
        set_cell_border(vc)
        rr = rc.paragraphs[0].add_run(k)
        rr.font.name  = "Calibri"
        rr.font.size  = Pt(10)
        rr.bold       = True
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        vr = vc.paragraphs[0].add_run(v)
        vr.font.name  = "Calibri"
        vr.font.size  = Pt(10)
        vr.font.color.rgb = DEEP_SLATE

    # Page break after cover
    doc.add_page_break()


# ─── Markdown Parser & Renderer ───────────────────────────────────────────────

def render_md_to_docx(doc: Document, md_path: Path) -> None:
    """
    Parse the Markdown FSD and render into the Word document.
    Handles: # headings, ## sub-headings, tables, code blocks,
             bullet lists, numbered lists, inline bold/code, horizontal rules.
    """
    lines = md_path.read_text(encoding="utf-8").splitlines()

    in_code_block   = False
    code_lines: list[str] = []
    in_table        = False
    table_rows: list[str] = []
    skip_cover      = True   # skip first 10 lines (YAML-like header already on cover page)
    line_idx        = 0

    while line_idx < len(lines):
        line = lines[line_idx]

        # ── Skip cover block (first metadata block) ──────────────────────
        if skip_cover and line_idx < 12:
            line_idx += 1
            continue
        skip_cover = False

        stripped = line.strip()

        # ── Code fence ───────────────────────────────────────────────────
        if stripped.startswith("```"):
            if in_code_block:
                # Close block
                in_code_block = False
                add_code_block(doc, code_lines)
                code_lines = []
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            else:
                in_code_block = True
                # Flush any pending table
                if in_table:
                    add_table_from_md(doc, table_rows)
                    table_rows = []
                    in_table   = False
            line_idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            line_idx += 1
            continue

        # ── Table rows ───────────────────────────────────────────────────
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_rows.append(stripped)
            line_idx += 1
            continue
        elif in_table:
            add_table_from_md(doc, table_rows)
            table_rows = []
            in_table   = False
            # Don't advance — re-process current line

        # ── Horizontal rules ─────────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            add_horizontal_rule(doc)
            line_idx += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if h_match:
            level   = len(h_match.group(1))
            h_text  = h_match.group(2).strip()
            # Strip markdown links
            h_text  = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', h_text)
            # Strip inline code markers
            h_text  = h_text.replace("`", "")
            style   = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}.get(level, "Heading 4")
            p       = doc.add_paragraph(h_text, style=style)
            if level == 1:
                add_horizontal_rule(doc)
            line_idx += 1
            continue

        # ── Bullet lists ─────────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text   = bullet_match.group(2).strip()
            # Strip markdown links
            text   = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
            add_bullet(doc, text, level=indent)
            line_idx += 1
            continue

        # ── Numbered lists ────────────────────────────────────────────────
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_match:
            indent = len(num_match.group(1)) // 2
            text   = num_match.group(2).strip()
            text   = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
            p      = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.25)
            pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
            parts   = pattern.split(text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2]); r.bold = True; r.font.color.rgb = DEEP_SLATE
                elif part.startswith("`") and part.endswith("`"):
                    r = p.add_run(part[1:-1]); r.font.name = "Courier New"; r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                else:
                    r = p.add_run(part); r.font.color.rgb = LIGHT_SLATE
                r.font.size = Pt(10.5)
            line_idx += 1
            continue

        # ── Block quote / indented formula lines ($$...$$) ────────────────
        if stripped.startswith("$$") or stripped == "$$":
            # Collect multi-line KaTeX formula and render as code block
            formula_lines = []
            if stripped != "$$":
                formula_lines.append(stripped)
                line_idx += 1
            else:
                line_idx += 1
                while line_idx < len(lines) and lines[line_idx].strip() != "$$":
                    formula_lines.append(lines[line_idx])
                    line_idx += 1
                line_idx += 1  # skip closing $$
            add_code_block(doc, formula_lines)
            line_idx += 1
            continue

        # ── Inline formula $...$ — render as body text ────────────────────
        if stripped.startswith("$") and not stripped.startswith("$$"):
            add_body_paragraph(doc, stripped, indent=1)
            line_idx += 1
            continue

        # ── Empty lines ───────────────────────────────────────────────────
        if not stripped:
            line_idx += 1
            continue

        # ── Body text ─────────────────────────────────────────────────────
        clean = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', stripped)
        add_body_paragraph(doc, clean)
        line_idx += 1

    # Flush any remaining table
    if in_table:
        add_table_from_md(doc, table_rows)
    if in_code_block:
        add_code_block(doc, code_lines)


# ─── Header / Footer ──────────────────────────────────────────────────────────

def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # Header
        hdr = section.header
        hp  = hdr.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = hp.add_run("PROMETHEUS  |  Functional Specification Document  |  CONFIDENTIAL")
        r1.font.name  = "Calibri"
        r1.font.size  = Pt(8)
        r1.font.color.rgb = LIGHT_SLATE
        # Header bottom border
        pPr  = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "D0D4D9")
        pBdr.append(bot)
        pPr.append(pBdr)

        # Footer
        ftr = section.footer
        fp  = ftr.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = fp.add_run("PROMETHEUS-FSD-v1.0  ·  Internal Use — Confidential  ·  April 5, 2026")
        r2.font.size      = Pt(8)
        r2.font.color.rgb = LIGHT_SLATE


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("Building PROMETHEUS FSD Word document...")

    doc = setup_document()
    add_cover_page(doc)
    render_md_to_docx(doc, SOURCE_MD)
    add_header_footer(doc)

    doc.save(str(OUTPUT_DOCX))
    print(f"\n✅ Document saved: {OUTPUT_DOCX}")
    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"   File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()

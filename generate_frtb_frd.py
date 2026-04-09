"""
FRTB Functional Requirements Document Generator

Generates a comprehensive DOCX document covering:
- Regulatory alignment with Basel MAR framework
- Business objectives and policy requirements
- Parameter specifications and market risk context
- Process workflows and technical specifications
- Glossary and equations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

def add_heading(doc, text, level=1):
    """Add a formatted heading."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_table_from_data(doc, headers, data, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
    
    return table

def generate_frtb_frd():
    """Generate the complete FRTB Functional Requirements Document."""
    
    doc = Document()
    
    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    
    title = doc.add_heading('PROMETHEUS Risk Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Fundamental Review of the Trading Book (FRTB)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.bold = True
    
    subtitle2 = doc.add_paragraph('Functional Requirements Document')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Document metadata
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = 'Light List Accent 1'
    
    meta_data = [
        ('Document Version:', '1.0'),
        ('Date:', date.today().strftime('%B %d, %Y')),
        ('Regulatory Framework:', 'Basel III MAR10-33'),
        ('Implementation:', 'Python 3.10+'),
        ('Status:', 'Production'),
        ('Classification:', 'Internal - Risk Management'),
    ]
    
    for i, (key, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        meta_table.rows[i].cells[1].text = value
    
    doc.add_page_break()
    
    # ========================================================================
    # TABLE OF CONTENTS (placeholder - Word will auto-generate)
    # ========================================================================
    
    add_heading(doc, 'Table of Contents', 1)
    doc.add_paragraph('[Auto-generated in Microsoft Word via References > Table of Contents]')
    doc.add_page_break()
    
    # ========================================================================
    # 1. EXECUTIVE SUMMARY
    # ========================================================================
    
    add_heading(doc, '1. Executive Summary', 1)
    
    doc.add_paragraph(
        'The FRTB (Fundamental Review of the Trading Book) engine implements the Basel Committee '
        'on Banking Supervision\'s comprehensive framework for market risk capital requirements, '
        'as specified in MAR10-33. This system replaces the previous market risk framework (Basel 2.5) '
        'with a more risk-sensitive approach that addresses shortcomings identified during the '
        '2007-2009 financial crisis.'
    )
    
    add_heading(doc, '1.1 Purpose and Scope', 2)
    
    doc.add_paragraph(
        'This functional requirements document specifies the business logic, regulatory alignment, '
        'and technical implementation of the FRTB capital calculation engine. The system calculates '
        'market risk capital using two methodologies:'
    )
    
    purposes = doc.add_paragraph()
    purposes.add_run('1. Sensitivities-Based Method (SBM)').bold = True
    purposes.add_run(' - Standardized approach based on regulatory risk weights and sensitivities (MAR21-23)\n')
    purposes.add_run('2. Internal Models Approach (IMA)').bold = True
    purposes.add_run(' - Bank-specific models using Expected Shortfall at 97.5% confidence (MAR31-33)')
    
    add_heading(doc, '1.2 Regulatory Mandate', 2)
    
    doc.add_paragraph(
        'The FRTB framework was finalized by the Basel Committee in January 2016 (revised January 2019) '
        'and is being implemented globally with staggered deadlines. Key regulatory objectives include:'
    )
    
    objectives = [
        'Replace Value-at-Risk (VaR) with Expected Shortfall (ES) for tail risk sensitivity',
        'Introduce a more granular and risk-sensitive standardized approach',
        'Strengthen the boundary between banking book and trading book',
        'Address model risk through eligibility tests and capital add-ons',
        'Incorporate liquidity horizons to reflect realistic exit timeframes',
        'Implement robust backtesting and profit-and-loss attribution frameworks',
    ]
    
    for obj in objectives:
        doc.add_paragraph(f'• {obj}', style='List Bullet')
    
    add_heading(doc, '1.3 System Capabilities', 2)
    
    capabilities_data = [
        ('Delta Risk', 'First-order price sensitivities across all risk classes'),
        ('Vega Risk', 'Volatility sensitivities for options and structured products'),
        ('Curvature Risk', 'Second-order (gamma) effects and price non-linearity'),
        ('Default Risk Charge', 'Jump-to-default and credit migration risk (MAR22)'),
        ('Residual Risk Add-On', 'Exotic risks not captured in sensitivities (MAR23)'),
        ('Expected Shortfall', 'Tail risk measure at 97.5% confidence (MAR31-33)'),
        ('Backtesting', 'Traffic-light framework for model validation (MAR99)'),
        ('Dynamic Adjustment', 'Real-time market condition integration'),
    ]
    
    add_table_from_data(
        doc,
        ['Component', 'Description'],
        capabilities_data,
        col_widths=[2, 4.5]
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 2. REGULATORY FRAMEWORK ALIGNMENT
    # ========================================================================
    
    add_heading(doc, '2. Regulatory Framework Alignment', 1)
    
    add_heading(doc, '2.1 Basel III Market Risk Framework (MAR)', 2)
    
    doc.add_paragraph(
        'The implementation adheres to the Basel Committee\'s "Minimum capital requirements for '
        'market risk" (MAR) standards, specifically:'
    )
    
    mar_standards = [
        ('MAR10', 'Boundary between banking book and trading book'),
        ('MAR20', 'Definitions and application (scope, risk classes)'),
        ('MAR21', 'Sensitivities-Based Method - Delta and Vega risks'),
        ('MAR22', 'Default Risk Charge (DRC)'),
        ('MAR23', 'Residual Risk Add-On (RRAO)'),
        ('MAR31', 'Internal Models Approach - Expected Shortfall'),
        ('MAR32', 'Profit and Loss Attribution test'),
        ('MAR33', 'Backtesting and capital floors'),
        ('MAR99', 'Supervisory reporting'),
    ]
    
    for standard, desc in mar_standards:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(standard).bold = True
        p.add_run(f': {desc}')
    
    add_heading(doc, '2.2 Implementation Status', 2)
    
    doc.add_paragraph(
        'The current implementation includes all regulatory fixes identified in the April 2026 '
        'code review. The following corrections ensure full MAR compliance:'
    )
    
    fixes_data = [
        ('FIX-01', 'HIGH', 'MAR21.6', 'Three-rho correlation formula correction'),
        ('FIX-02', 'HIGH', 'MAR23.6', 'Curvature charge sqrt() removal'),
        ('FIX-03', 'MEDIUM', 'MAR21.44', 'GIRR risk weight count (13→10 values)'),
        ('FIX-04', 'MEDIUM', 'MAR33.8', 'ES tail count (int→ceil)'),
        ('FIX-05', 'MEDIUM', 'MAR21.73', 'CSR_SEC/CTP parameters added'),
        ('FIX-06', 'LOW', 'MAR31.14', 'NMRF charge default (0.0015→0.0)'),
    ]
    
    add_table_from_data(
        doc,
        ['Fix ID', 'Severity', 'MAR Ref', 'Description'],
        fixes_data,
        col_widths=[1, 1, 1.2, 3.3]
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 3. BUSINESS AND POLICY OBJECTIVES
    # ========================================================================
    
    add_heading(doc, '3. Business and Policy Objectives', 1)
    
    add_heading(doc, '3.1 Capital Management', 2)
    
    doc.add_paragraph(
        'The FRTB engine serves as the primary calculation tool for market risk Pillar 1 capital. '
        'Key business objectives include:'
    )
    
    capital_objectives = [
        ('Regulatory Capital', 'Calculate minimum required capital per MAR21-33 for regulatory reporting'),
        ('Economic Capital', 'Support internal capital allocation and risk-adjusted performance measurement'),
        ('Stress Testing', 'Enable scenario analysis and stress capital buffer determination'),
        ('Risk Limits', 'Inform trading desk limits and risk appetite framework'),
        ('Model Governance', 'Facilitate model validation, backtesting, and P&L attribution'),
    ]
    
    for title, desc in capital_objectives:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)
    
    add_heading(doc, '3.2 Risk Measurement Philosophy', 2)
    
    doc.add_paragraph(
        'The system implements a dual approach to market risk measurement, balancing '
        'standardization with model sophistication:'
    )
    
    p1 = doc.add_paragraph()
    p1.add_run('Sensitivities-Based Method (SBM)').bold = True
    p1.add_run(
        '\nProvides a transparent, comparable, and verifiable standardized measure. Uses regulatory '
        'risk weights and correlation assumptions. Suitable for less complex portfolios and serves '
        'as a floor for IMA capital.'
    )
    
    p2 = doc.add_paragraph()
    p2.add_run('Internal Models Approach (IMA)').bold = True
    p2.add_run(
        '\nReflects bank-specific risk profiles through historical simulation or parametric methods. '
        'Requires regulatory approval and ongoing validation. Allows capital efficiency for well-modeled risks.'
    )
    
    add_heading(doc, '3.3 Trading Desk Coverage', 2)
    
    doc.add_paragraph(
        'The engine supports all major asset classes and trading strategies defined in MAR20:'
    )
    
    coverage_data = [
        ('Interest Rate (GIRR)', 'Government bonds, swaps, swaptions, caps/floors, inflation products'),
        ('Credit Spread (CSR)', 'Corporate bonds, CDS, securitizations, tranches, CDO, CLO'),
        ('Equity', 'Cash equities, equity derivatives, index products, convertibles'),
        ('Foreign Exchange (FX)', 'Spot, forwards, FX options, cross-currency swaps'),
        ('Commodity (CMDTY)', 'Energy, metals, agriculture, freight, weather derivatives'),
    ]
    
    add_table_from_data(
        doc,
        ['Risk Class', 'Instruments Covered'],
        coverage_data,
        col_widths=[2, 4.5]
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 4. SYSTEM ARCHITECTURE
    # ========================================================================
    
    add_heading(doc, '4. System Architecture', 1)
    
    add_heading(doc, '4.1 High-Level Design', 2)
    
    doc.add_paragraph(
        'The FRTB engine follows a modular, layered architecture that separates data acquisition, '
        'business logic, and result aggregation:'
    )
    
    doc.add_paragraph()
    
    # ASCII architecture diagram
    architecture_diagram = """
    ┌─────────────────────────────────────────────────────────────────┐
    │                      FRTB Engine (FRTBEngine)                   │
    │                                                                 │
    │  Input: Sensitivities, P&L History, DRC Positions              │
    │  Output: FRTBResult (SBM, IMA, DRC, RRAO, Capital)             │
    └────────────┬────────────────────────────────────┬───────────────┘
                 │                                    │
                 ▼                                    ▼
    ┌────────────────────────┐           ┌────────────────────────────┐
    │   SBM Calculator       │           │    IMA Calculator          │
    │   (SBMCalculator)      │           │    (IMACalculator)         │
    │                        │           │                            │
    │ • Delta Risk           │           │ • Historical ES            │
    │ • Vega Risk            │           │ • Parametric ES            │
    │ • Curvature Risk       │           │ • Stressed ES              │
    │ • Correlation Model    │           │ • Liquidity Horizons       │
    └────────────┬───────────┘           └────────────┬───────────────┘
                 │                                    │
                 ▼                                    ▼
    ┌────────────────────────────────────────────────────────────────┐
    │              Market Data & Configuration Layer                 │
    │                                                                │
    │  • MarketDataFeed (real-time, Bloomberg/FRED/yfinance)        │
    │  • FRTBConfig (risk weights, correlations, parameters)        │
    │  • DynamicParameterAdjustment (stress-conditional scaling)    │
    └────────────────────────────────────────────────────────────────┘
    """
    
    arch_para = doc.add_paragraph(architecture_diagram, style='No Spacing')
    arch_para.runs[0].font.name = 'Courier New'
    arch_para.runs[0].font.size = Pt(9)
    
    add_heading(doc, '4.2 Core Components', 2)
    
    components_data = [
        ('FRTBEngine', 'Orchestrates calculation, validates inputs, manages configuration, routes to SBM/IMA'),
        ('SBMCalculator', 'Implements MAR21-23: delta/vega/curvature charges, correlation aggregation, DRC, RRAO'),
        ('IMACalculator', 'Implements MAR31-33: ES calculation, liquidity horizon scaling, backtesting, P&L attribution'),
        ('MarketDataFeed', 'Fetches real-time market conditions (VIX, spreads, volatilities) from external sources'),
        ('FRTBConfig', 'Stores risk weights, correlations, thresholds, regulatory parameters per MAR21'),
        ('CorrelationModel', 'Computes intra-bucket and inter-bucket correlation matrices (MAR21.58-60)'),
        ('DynamicParameterAdjustment', 'Adjusts risk weights/correlations based on market regime (Pillar 2)'),
    ]
    
    add_table_from_data(
        doc,
        ['Component', 'Responsibility'],
        components_data,
        col_widths=[2, 4.5]
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 5. PARAMETER SPECIFICATIONS
    # ========================================================================
    
    add_heading(doc, '5. Parameter Specifications', 1)
    
    add_heading(doc, '5.1 Market Conditions Parameters', 2)
    
    doc.add_paragraph(
        'The MarketConditions dataclass captures real-time market state for dynamic risk weight '
        'adjustment and regime classification. These parameters are sourced externally via the '
        'MarketDataFeed three-tier fallback system (Bloomberg → FRED/yfinance → conservative defaults).'
    )
    
    doc.add_paragraph()
    
    market_params_data = [
        ('vix_level', 'float', 'CBOE VIX Index level (index points)', 'VIXCLS (FRED) / ^VIX (yfinance)', 
         'Fear gauge; normal ~15-20, crisis >40', '20.0'),
        ('equity_vol_index', 'float', 'Equity realized volatility (% annualized)', 'VXEEMCLS (FRED) / ^VIX', 
         'Equity market turbulence indicator', '22.0%'),
        ('credit_spread_ig', 'float', 'Investment Grade credit spread (bp)', 'BAMLC0A0CMEY (FRED)', 
         'Corporate credit health; normal ~100bp', '120 bp'),
        ('credit_spread_hy', 'float', 'High Yield credit spread (bp)', 'BAMLH0A0HYM2 (FRED)', 
         'Distressed credit indicator; normal ~400bp', '450 bp'),
        ('fx_vol_index', 'float', 'FX volatility index (% annualized)', 'DEXJPUS realized vol (FRED)', 
         'Currency market stress; normal ~8%', '9.0%'),
        ('cmdty_vol_index', 'float', 'Commodity volatility (% annualized)', 'CL=F realized vol (yfinance)', 
         'Commodity market turbulence; oil proxy', '28.0%'),
        ('ir_vol_swaption', 'float', 'Interest rate volatility (bp)', 'MOVE Index (FRED)', 
         'Swaption implied vol; normal ~50bp', '60 bp'),
    ]
    
    # Create detailed table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    headers = ['Parameter', 'Type', 'Description', 'Data Source', 'Market Risk Context', 'Fallback']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    
    for param_data in market_params_data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(param_data):
            row_cells[i].text = str(cell_data)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
    
    doc.add_paragraph()
    
    add_heading(doc, '5.1.1 Why These Parameters?', 3)
    
    doc.add_paragraph(
        'Each parameter serves a specific purpose in market risk assessment and aligns with '
        'established academic research and industry practice:'
    )
    
    why_params = [
        ('VIX Level', 'Granger-causal relationship with equity returns (Whaley 2000); '
         'forward-looking fear gauge that predicts realized volatility'),
        ('Credit Spreads', 'Leading indicator of economic stress (Gilchrist & Zakrajšek 2012); '
         'HY-IG differential captures credit cycle positioning'),
        ('Volatility Indices', 'Direct input to option pricing and Greeks; correlates with '
         'liquidity conditions and market depth (Amihud illiquidity measure)'),
        ('MOVE Index', 'Fixed income volatility benchmark; predicts interest rate risk and '
         'duration hedging costs'),
    ]
    
    for param, rationale in why_params:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(f'{param}: ').bold = True
        p.add_run(rationale)
    
    add_heading(doc, '5.1.2 How Are They Used?', 3)
    
    doc.add_paragraph(
        'Parameters feed into three key business processes:'
    )
    
    doc.add_paragraph(
        '1. Market Regime Classification\n'
        '   • Composite stress index = 0.40×VIX_norm + 0.35×Spread_norm + 0.25×EqVol_norm\n'
        '   • Normal: stress < 30% | Stressed: 30-65% | Crisis: >65%\n'
        '   • Used for scenario selection and management reporting',
        style='List Number'
    )
    
    doc.add_paragraph(
        '2. Dynamic Risk Weight Adjustment (Pillar 2 / ICAAP)\n'
        '   • Scaling factor = base_multiplier + vol_premium × (realized_vol / normal_vol)\n'
        '   • ONLY applied to internal stress capital; regulatory SBM uses prescribed MAR21 weights\n'
        '   • Example: EQ risk weight scales 0.75-1.50× based on equity_vol_index',
        style='List Number'
    )
    
    doc.add_paragraph(
        '3. Correlation Adjustment\n'
        '   • Crisis conditions: correlations increase (diversification breakdown)\n'
        '   • Multiplier = 1.0 (normal) to 1.50 (crisis)\n'
        '   • Reflects empirical evidence from 2008 crisis (Forbes & Rigobon 2002)',
        style='List Number'
    )
    
    add_heading(doc, '5.1.3 What Do They Achieve?', 3)
    
    doc.add_paragraph(
        'The real-time market data integration achieves three strategic objectives:'
    )
    
    achievements = [
        ('Counter-Cyclical Capital', 'Higher capital in stressed markets prevents pro-cyclical '
         'deleveraging; aligns with Basel III counter-cyclical buffer framework'),
        ('Early Warning System', 'Stress index crossing thresholds triggers risk committee alerts '
         'and enhanced monitoring protocols'),
        ('Model Risk Mitigation', 'External validation of internal risk estimates; flags when model '
         'parameters diverge from market reality'),
    ]
    
    for title, desc in achievements:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ========================================================================
    # 5.2 SBM RISK WEIGHTS
    # ========================================================================
    
    add_heading(doc, '5.2 Sensitivities-Based Method (SBM) Risk Weights', 2)
    
    doc.add_paragraph(
        'Risk weights (RW) are regulatory parameters prescribed in MAR21.44-96. They represent '
        'the expected percentage change in portfolio value for a standardized shock to each risk factor. '
        'All values are stored in FRTBConfig.delta_rw as floating-point fractions (e.g., 0.017 = 1.7%).'
    )
    
    add_heading(doc, '5.2.1 Interest Rate Risk Weights (GIRR)', 3)
    
    doc.add_paragraph(
        'MAR21.44 Table 2 prescribes exactly 10 risk weights corresponding to the regulatory tenor grid: '
        '3m, 6m, 1y, 2y, 3y, 5y, 10y, 15y, 20y, 30y. These apply to all currencies.'
    )
    
    girr_rw_data = [
        ('3m', '0.017', '1.7%', 'Short-dated instruments; low duration sensitivity'),
        ('6m', '0.017', '1.7%', 'Money market tenors'),
        ('1y', '0.016', '1.6%', 'Standard benchmark tenor'),
        ('2y', '0.013', '1.3%', 'Decreasing RW reflects lower yield vol at intermediate tenors'),
        ('3y', '0.012', '1.2%', ''),
        ('5y', '0.011', '1.1%', 'Liquid benchmark tenor (5y swap reference)'),
        ('10y', '0.011', '1.1%', 'Key benchmark; most liquid government bonds'),
        ('15y', '0.011', '1.1%', 'Long-dated benchmarks'),
        ('20y', '0.011', '1.1%', ''),
        ('30y', '0.011', '1.1%', 'Ultra-long tenor; duration risk'),
    ]
    
    add_table_from_data(
        doc,
        ['Tenor', 'RW (decimal)', 'RW (%)', 'Rationale'],
        girr_rw_data,
        col_widths=[0.8, 1.2, 1, 3.5]
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Why These Values? ').bold = True
    p.add_run(
        'Risk weights reflect historical yield volatility calibrated to stressed periods '
        '(2007-2009 crisis + 2011 sovereign debt crisis). The hump shape (higher at short end, '
        'decreasing to 2-3y, flat thereafter) matches empirical yield curve dynamics.'
    )
    
    p2 = doc.add_paragraph()
    p2.add_run('Special Cases (MAR21.52): ').bold = True
    
    special_cases = doc.add_paragraph()
    special_cases.add_run('• Inflation: ').bold = True
    special_cases.add_run('1.6% (same as 10y nominal)\n')
    special_cases.add_run('• Cross-Currency Basis: ').bold = True
    special_cases.add_run('1.6% (supervisory proxy; limited historical data)')
    
    add_heading(doc, '5.2.2 Credit Spread Risk Weights (CSR)', 3)
    
    doc.add_paragraph(
        'Credit risk weights vary by sector, rating, and securitization type per MAR21.67-76.'
    )
    
    csr_data = [
        ('Sovereigns (AAA-AA)', '0.5%', 'Low default risk; liquid benchmarks'),
        ('Sovereigns (A)', '1.0%', ''),
        ('Sovereigns (BBB)', '3.0%', 'Emerging market exposure'),
        ('Corporates (IG)', '0.5-2.0%', 'Varies by sector; Financials 0.5%, Consumers 1.0%'),
        ('Corporates (HY)', '3.0-14.0%', 'BB rated 3%, B rated 8%, CCC+ 14%'),
        ('Securitization (Senior IG)', '0.9%', 'MAR21.73: RMBS/CMBS senior tranches'),
        ('Securitization (Non-Senior)', '1.5-2.5%', 'Mezzanine and first-loss pieces'),
        ('CTP Indices', '1.6-4.0%', 'CDX/iTraxx index tranches; correlation trading'),
    ]
    
    add_table_from_data(
        doc,
        ['Category', 'Risk Weight', 'Notes'],
        csr_data,
        col_widths=[2.5, 1.5, 2.5]
    )
    
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run('Rationale: ').bold = True
    p3.add_run(
        'Weights increase with credit risk (rating deterioration) and structural subordination. '
        'HY weights are 3-28× higher than IG, reflecting higher default probability and loss-given-default. '
        'Securitization weights incorporate model risk premium (complex waterfall structures).'
    )
    
    add_heading(doc, '5.2.3 Equity Risk Weights', 3)
    
    doc.add_paragraph('MAR21.78-82 distinguishes large-cap (>$2B market cap) from small-cap by sector.')
    
    eq_data = [
        ('Emerging Markets', 'Large: 55%', 'Small: 70%', 'Higher volatility; FX and political risk'),
        ('Developed - Cyclicals', 'Large: 60%', 'Small: 70%', 'Consumer discretionary, industrials'),
        ('Developed - Defensives', 'Large: 45%', 'Small: 70%', 'Utilities, healthcare, consumer staples'),
        ('Financials', 'Large: 55%', 'Small: 70%', 'Banks, insurance, real estate'),
        ('Technology', 'Large: 30%', 'Small: 70%', 'Lower systematic risk; sector-specific vol'),
        ('Other', 'Large: 50%', 'Small: 70%', 'Catch-all bucket'),
    ]
    
    add_table_from_data(
        doc,
        ['Sector', 'Large-Cap RW', 'Small-Cap RW', 'Rationale'],
        eq_data,
        col_widths=[2, 1.3, 1.3, 2]
    )
    
    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.add_run('Why Large-Cap vs. Small-Cap? ').bold = True
    p4.add_run(
        'Small-cap equities exhibit 1.4-1.5× higher realized volatility and lower liquidity. '
        'Uniform 70% RW reflects limited historical differentiation by sector for small-caps. '
        'Large-cap sector granularity captures beta differences (tech ~0.9 vs. financials ~1.3).'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 5.3 CORRELATION PARAMETERS
    # ========================================================================
    
    add_heading(doc, '5.3 Correlation Parameters', 2)
    
    doc.add_paragraph(
        'Correlations govern diversification benefits across risk factors. MAR21.58-60 (intra-bucket) '
        'and MAR21.61-96 (inter-bucket) provide prescribed values. The CorrelationModel class '
        'implements these formulas.'
    )
    
    add_heading(doc, '5.3.1 Intra-Bucket Correlations', 3)
    
    doc.add_paragraph(
        'Intra-bucket ρ applies to sensitivities within the same regulatory bucket (e.g., same currency, '
        'same sector). Higher values mean stronger diversification offset.'
    )
    
    intra_corr_data = [
        ('GIRR (nominal)', '99.9%', 'MAR21.58: ρ(t₁,t₂) = exp(-0.03×|ln(t₁/t₂)|)', 
         'Near-perfect correlation across yield curve; parallel shifts dominate'),
        ('GIRR (inflation)', '40%', 'MAR21.60: Inflation vs. nominal within same currency', 
         'Decoupling during QE periods'),
        ('CSR Non-Sec', '65%', 'MAR21.71: Same sector and rating bucket', 
         'Idiosyncratic credit risk reduces correlation'),
        ('CSR Securitization', '99%', 'MAR21.73: Senior tranches of same collateral pool', 
         'Common underlying asset pool drives correlation'),
        ('Equity Large-Cap', '15%', 'MAR21.80: Same sector, different issuers', 
         'Company-specific factors dominate'),
        ('Equity Small-Cap', '7.5%', 'MAR21.80: Lower liquidity, higher idiosyncratic risk', 
         'Less co-movement than large-cap'),
        ('FX', '100%', 'MAR21.90: All FX pairs within bucket', 
         'Supervisory simplification; triangular arbitrage'),
        ('Commodity', '55%', 'MAR21.94: Same commodity type (e.g., energy, metals)', 
         'Supply-demand shocks are commodity-specific'),
    ]
    
    add_table_from_data(
        doc,
        ['Risk Class', 'ρ (intra)', 'Formula / Reference', 'Economic Rationale'],
        intra_corr_data
    )
    
    doc.add_paragraph()
    
    add_heading(doc, '5.3.2 Inter-Bucket Correlations (γ)', 3)
    
    doc.add_paragraph(
        'Inter-bucket γ applies across different buckets (e.g., different currencies, sectors, asset classes). '
        'Lower than intra-bucket correlations to reflect structural differences.'
    )
    
    inter_corr_data = [
        ('GIRR', '50%', 'MAR21.60: Cross-currency correlation', 
         'Central bank policy divergence, but global rate cycle linkage'),
        ('CSR Non-Sec', '0-50%', 'MAR21.76: Sector-dependent; same-sector IG 35%, cross-sector can be 0%', 
         'Corporate correlation matrix based on empirical equity correlation'),
        ('Equity', '15%', 'MAR21.82: Cross-sector within same market', 
         'Market factor explains ~15-20% of variance (Fama-French)'),
        ('FX', '60%', 'MAR21.90: Different currency pairs', 
         'USD common factor; EM pairs show higher correlation'),
        ('Commodity', '20%', 'MAR21.96: Different commodity types (energy vs. metals)', 
         'Low correlation; different supply chains'),
    ]
    
    add_table_from_data(
        doc,
        ['Risk Class', 'γ (inter)', 'Reference', 'Rationale'],
        inter_corr_data
    )
    
    doc.add_paragraph()
    
    p5 = doc.add_paragraph()
    p5.add_run('Critical Note: ').bold = True
    p5.add_run(
        'MAR21.60 specifies that GIRR_INFLATION and GIRR_XCCY_BASIS receive γ=0 against nominal GIRR '
        'curves within the same currency. This sub-curve zero-correlation cannot be represented by a '
        'single scalar. The implementation treats these as separate risk_class entries to preserve '
        'regulatory accuracy.'
    )
    
    add_heading(doc, '5.3.3 Three-Scenario Aggregation (MAR21.6)', 3)
    
    doc.add_paragraph(
        'To capture tail correlation behavior, MAR21.6 requires calculating capital under three '
        'correlation scenarios and taking the maximum:'
    )
    
    doc.add_paragraph('Scenario 1: Base correlations (ρ, γ as configured)')
    doc.add_paragraph('Scenario 2: High correlations - ρ_high = ρ + 0.25×(1-ρ), γ_high = γ + 0.25×(1-γ)')
    doc.add_paragraph('Scenario 3: Low correlations - ρ_low = max(2ρ-1, 0), γ_low = max(2γ-1, 0)')
    
    doc.add_paragraph()
    p6 = doc.add_paragraph()
    p6.add_run('Rationale: ').bold = True
    p6.add_run(
        'Correlation is unstable and model-dependent. The three scenarios ensure capital adequacy across '
        'plausible correlation regimes. High scenario captures contagion/crisis (correlation →1). '
        'Low scenario captures hedge effectiveness uncertainty. Maximum of three provides conservative buffer.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 6. PROCESS WORKFLOWS
    # ========================================================================
    
    add_heading(doc, '6. Process Workflows', 1)
    
    add_heading(doc, '6.1 Daily Capital Calculation Workflow', 2)
    
    doc.add_paragraph('Standard daily production run for regulatory reporting (MAR99).')
    
    doc.add_paragraph()
    
    workflow_text = """
    1. DATA INGESTION
       ├─ Extract trade sensitivities from front-office systems (delta, vega, gamma)
       ├─ Retrieve 250-day P&L history from data warehouse
       ├─ Fetch current market data (MarketDataFeed.get_current_conditions())
       │  └─ Three-tier fetch: Bloomberg → FRED → yfinance → fallback
       └─ Load DRC positions (credit names, notionals, ratings, maturities)
    
    2. DATA VALIDATION
       ├─ Check sensitivity completeness (all risk factors mapped to buckets)
       ├─ Validate P&L series (no gaps, outliers flagged)
       ├─ Verify market data freshness (timestamp < 24h)
       └─ Confirm desk eligibility for IMA (approved trading desks only)
    
    3. SBM CALCULATION (MAR21-23)
       ├─ Aggregate sensitivities by risk class and bucket
       ├─ Apply risk weights from FRTBConfig.delta_rw
       ├─ Compute weighted sensitivities (WS = sensitivity × RW)
       ├─ Build correlation matrices (intra-bucket ρ, inter-bucket γ)
       ├─ Three-scenario aggregation:
       │  ├─ Scenario 1 (base): K_base = √(WS' × Σ_base × WS)
       │  ├─ Scenario 2 (high): K_high = √(WS' × Σ_high × WS)
       │  └─ Scenario 3 (low):  K_low  = √(WS' × Σ_low  × WS)
       ├─ Delta charge = max(K_base, K_high, K_low)
       ├─ Vega charge (same methodology for vega sensitivities)
       ├─ Curvature charge (CVR = max(CVR_up, CVR_dn) per MAR21.8)
       ├─ DRC (jump-to-default + migration, MAR22)
       ├─ RRAO (exotic/residual risks, MAR23.5)
       └─ Total SBM = Delta + Vega + Curvature + DRC + RRAO
    
    4. IMA CALCULATION (MAR31-33)
       ├─ Filter P&L to eligible desk scope
       ├─ Compute Expected Shortfall at 97.5% confidence:
       │  ├─ Sort P&L in ascending order
       │  ├─ Tail index = ceil(0.025 × n_observations)
       │  └─ ES = -mean(P&L[1:tail_index])  [note: negative of losses]
       ├─ Liquidity horizon adjustments (MAR33.4):
       │  └─ ES_LH = ES × √(LH_days / 10)  [separate ES per risk class, then aggregate]
       ├─ Stressed ES (P&L from 12-month stressed period, MAR33.6):
       │  └─ Stressed_ES = max(ES_current_data, ES_stressed_period)
       ├─ Apply IMA multiplier (MAR33.9):
       │  └─ mc = 1.5 + addon(backtesting_exceptions)  [1.5 to 2.0 range]
       └─ IMA Capital = mc × max(ES_LH, Stressed_ES)
    
    5. CAPITAL DETERMINATION
       ├─ Regulatory capital = max(SBM, IMA)  [MAR33.8]
       ├─ Apply capital floor: capital ≥ 0.72 × SBM  [MAR99 transitional floor]
       └─ Aggregate across trading desks
    
    6. OUTPUT & REPORTING
       ├─ Generate FRTBResult object with full breakdown
       ├─ Store to database (audit trail, time series)
       ├─ Publish to risk dashboard
       ├─ Trigger alerts if limits breached
       └─ Prepare regulatory templates (MAR99 reporting)
    """
    
    doc.add_paragraph(workflow_text, style='No Spacing').runs[0].font.name = 'Courier New'
    doc.add_paragraph(workflow_text, style='No Spacing').runs[0].font.size = Pt(9)
    
    doc.add_page_break()
    
    add_heading(doc, '6.2 Market Data Refresh Workflow', 2)
    
    doc.add_paragraph(
        'The MarketDataFeed implements a three-tier fallback strategy with caching. '
        'This workflow ensures data is always available even during source outages.'
    )
    
    mdf_workflow = """
    MarketDataFeed.get_current_conditions(force_refresh=False)
    │
    ├─ Check in-memory cache
    │  └─ If valid (age < TTL): return cached MarketConditions ✓
    │
    ├─ TIER 1: Bloomberg / Refinitiv (if configured)
    │  ├─ Try Bloomberg BSAPI (blpapi.Session)
    │  │  └─ Tickers: VIX Index, LF98OAS Index, MOVE Index, etc.
    │  ├─ Try Refinitiv Data API (refinitiv.data)
    │  │  └─ RICs: .VIX, .MOVE, etc.
    │  └─ If successful: transform & cache → return ✓
    │
    ├─ TIER 2: FRED / yfinance (public sources)
    │  ├─ Try FRED JSON API (if FRED_API_KEY set)
    │  │  └─ Series: VIXCLS, BAMLH0A0HYM2, MOVE, etc.
    │  ├─ Fallback to FRED CSV endpoint (no auth)
    │  │  └─ URL: https://fred.stlouisfed.org/graph/fredgraph.csv?id={series}
    │  ├─ Try yfinance (for VIX, commodity vol)
    │  │  └─ Tickers: ^VIX, CL=F (with realized vol calculation)
    │  └─ If successful: transform & cache → return ✓
    │
    ├─ TIER 3: Hardcoded Conservative Fallbacks
    │  └─ Return MarketConditions with _FALLBACK values
    │     (slightly above historical medians → conservative for capital)
    │
    └─ Cache result (in-memory + optional disk JSON)
       ├─ In-memory: protected by threading.RLock
       └─ Disk: /tmp/frtb_market_cache.json (shareable across processes)
    """
    
    mdf_para = doc.add_paragraph(mdf_workflow, style='No Spacing')
    mdf_para.runs[0].font.name = 'Courier New'
    mdf_para.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    p7 = doc.add_paragraph()
    p7.add_run('Design Rationale: ').bold = True
    p7.add_run(
        'The three-tier design ensures 100% uptime. Even complete network failure produces valid '
        'output (conservative fallbacks). Tier ordering prioritizes accuracy (Bloomberg real-time) '
        'over cost (free FRED), with built-in redundancy.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 7. KEY EQUATIONS
    # ========================================================================
    
    add_heading(doc, '7. Key Equations and Formulas', 1)
    
    add_heading(doc, '7.1 SBM Delta Risk Aggregation (MAR21.3-6)', 2)
    
    doc.add_paragraph('For a single risk class with B buckets:')
    
    doc.add_paragraph('Step 1: Weighted Sensitivity for factor k in bucket b')
    doc.add_paragraph('    WS_bk = s_bk × RW_k')
    doc.add_paragraph('    where s_bk = net sensitivity, RW_k = risk weight')
    
    doc.add_paragraph()
    doc.add_paragraph('Step 2: Bucket-level capital charge (intra-bucket aggregation)')
    doc.add_paragraph('    K_b = √(Σ_k WS_bk² + Σ_k Σ_l≠k ρ_kl × WS_bk × WS_bl)')
    doc.add_paragraph('    where ρ_kl = intra-bucket correlation between factors k and l')
    
    doc.add_paragraph()
    doc.add_paragraph('Step 3: Risk class capital (inter-bucket aggregation)')
    doc.add_paragraph('    K = √(Σ_b K_b² + Σ_b Σ_c≠b γ_bc × S_b × S_c)')
    doc.add_paragraph('    where γ_bc = inter-bucket correlation, S_b = Σ_k WS_bk (bucket sensitivity sum)')
    
    doc.add_paragraph()
    doc.add_paragraph('Step 4: Three-scenario maximum (MAR21.6)')
    doc.add_paragraph('    Delta_risk_class = max(K_base, K_high, K_low)')
    doc.add_paragraph('    where each K uses adjusted correlations per scenario')
    
    doc.add_paragraph()
    p8 = doc.add_paragraph()
    p8.add_run('Implementation: ').bold = True
    p8.add_run(
        'SBMCalculator._aggregate_buckets() builds the full covariance matrix Σ with intra-bucket '
        'ρ on off-diagonals within each bucket, and inter-bucket γ across buckets. Three matrices '
        '(base, high, low) are constructed with correlation adjustments, then capital is computed '
        'as √(WS\' × Σ × WS) for each scenario.'
    )
    
    doc.add_page_break()
    
    add_heading(doc, '7.2 Curvature Risk (MAR21.8-10)', 2)
    
    doc.add_paragraph('Curvature captures convexity (gamma) risk not reflected in delta.')
    
    doc.add_paragraph('For each risk factor k:')
    doc.add_paragraph('    CVR_k = -min(Σ_i V_i(x_k + Δx_k) - V_i(x_k), Σ_i V_i(x_k - Δx_k) - V_i(x_k))')
    doc.add_paragraph('    where:')
    doc.add_paragraph('        V_i(x) = trade i value at risk factor level x')
    doc.add_paragraph('        Δx_k   = risk weight shock to factor k')
    doc.add_paragraph('        CVR_k  = worst-case loss (upside move vs. downside move)')
    
    doc.add_paragraph()
    doc.add_paragraph('Aggregation formula (MAR21.9):')
    doc.add_paragraph('    Ξ_b = Σ_k Ψ_k(CVR_k) + Σ_k Σ_l≠k ρ_kl × Ψ_k(CVR_k) × Ψ_l(CVR_l)')
    doc.add_paragraph('    where Ψ_k(x) = sign(x) × √|x|  (theta function for smooth aggregation)')
    
    doc.add_paragraph()
    doc.add_paragraph('Final curvature charge:')
    doc.add_paragraph('    Curvature = √(Σ_b Ξ_b² + Σ_b Σ_c≠b γ_bc × Ξ_b × Ξ_c)')
    
    doc.add_paragraph()
    p9 = doc.add_paragraph()
    p9.add_run('FIX-02 Note: ').bold = True
    p9.add_run(
        'Previous implementation erroneously applied a final sqrt() to the curvature charge. '
        'MAR23.6 specifies that Ξ_s (aggregated curvature) is already in dollar capital units. '
        'The extra sqrt() was removed in April 2026 code review.'
    )
    
    doc.add_page_break()
    
    add_heading(doc, '7.3 Expected Shortfall (MAR31.5)', 2)
    
    doc.add_paragraph('ES measures average loss in the tail beyond the 97.5% quantile.')
    
    doc.add_paragraph('Historical simulation approach:')
    doc.add_paragraph('1. Sort 1-day P&L history in ascending order (worst to best)')
    doc.add_paragraph('2. Tail index = ceil(2.5% × n_observations)')
    doc.add_paragraph('3. ES = -mean(P&L[1 : tail_index])')
    doc.add_paragraph('   (negative because P&L is stated as profit; ES is a loss measure)')
    
    doc.add_paragraph()
    doc.add_paragraph('Liquidity horizon adjustment (MAR33.4):')
    doc.add_paragraph('    ES_LH(risk_class) = ES_10day(risk_class) × √(LH_days / 10)')
    doc.add_paragraph('    where LH_days ∈ {10, 20, 40, 60, 120} per MAR33 Table 1')
    
    doc.add_paragraph()
    doc.add_paragraph('Aggregation across risk classes (MAR33.5):')
    doc.add_paragraph('    ES_total = √(Σ_rc ES_LH(rc)² + 2 × Σ_rc Σ_sc≠rc ρ_rc,sc × ES_LH(rc) × ES_LH(sc))')
    doc.add_paragraph('    where ρ_rc,sc = inter-risk-class correlation from IMCC calculation')
    
    doc.add_paragraph()
    doc.add_paragraph('Final IMA capital (MAR33.8):')
    doc.add_paragraph('    IMA = mc × max(ES_total, Stressed_ES, 60d_avg_IMCC)')
    doc.add_paragraph('    where mc = 1.5 + addon(backtesting_exceptions)  [1.5 to 2.0]')
    
    doc.add_paragraph()
    p10 = doc.add_paragraph()
    p10.add_run('FIX-04 Note: ').bold = True
    p10.add_run(
        'Tail index calculation changed from int(0.025 × n) to ceil(0.025 × n) per MAR33.8. '
        'Using int() systematically understated ES by excluding one tail observation. For n=250, '
        'this means 7 observations instead of 6, correcting a ~14% bias.'
    )
    
    doc.add_page_break()
    
    add_heading(doc, '7.4 Dynamic Risk Weight Adjustment (Pillar 2)', 2)
    
    doc.add_paragraph(
        'Dynamic adjustment is an INTERNAL tool for stress testing and ICAAP. It is NEVER applied '
        'to regulatory SBM capital (which uses prescribed MAR21 weights unchanged).'
    )
    
    doc.add_paragraph('Risk class-specific scaling formulas:')
    
    doc.add_paragraph()
    doc.add_paragraph('Equity risk class:')
    doc.add_paragraph('    scaling = 0.75 + 0.75 × (equity_vol_index / 20.0)')
    doc.add_paragraph('    Example: vol=30% → scaling=1.875× (severe stress)')
    doc.add_paragraph('            vol=15% → scaling=1.313× (moderate)')
    
    doc.add_paragraph()
    doc.add_paragraph('Credit risk class (CSR):')
    doc.add_paragraph('    scaling = 0.80 + 0.60 × (credit_spread_hy / 400.0)')
    doc.add_paragraph('    Example: HY=800bp → scaling=2.00× (crisis)')
    doc.add_paragraph('            HY=400bp → scaling=1.40× (normal)')
    
    doc.add_paragraph()
    doc.add_paragraph('Crisis amplification (if stress_index > 0.65):')
    doc.add_paragraph('    scaling ← scaling × (1.0 + 0.5 × (stress - 0.65) / 0.35)')
    doc.add_paragraph('    Capped at 3.0× to prevent unrealistic capital levels')
    
    doc.add_paragraph()
    p11 = doc.add_paragraph()
    p11.add_run('Rationale: ').bold = True
    p11.add_run(
        'Volatility regime shifts are mean-reverting but persistent (GARCH effects). '
        'Historical 1Y average vol has 60% correlation with forward 1Y realized vol (Christensen & Prabhala 1998). '
        'Scaling factors are calibrated to historical crisis episodes (1998 LTCM, 2008 Lehman, 2020 COVID).'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # 8. GLOSSARY
    # ========================================================================
    
    add_heading(doc, '8. Glossary of Terms', 1)
    
    glossary_data = [
        ('FRTB', 'Fundamental Review of the Trading Book', 'Basel Committee\'s revised market risk framework (2016/2019)'),
        ('SBM', 'Sensitivities-Based Method', 'Standardized approach using regulatory risk weights and sensitivities (MAR21)'),
        ('IMA', 'Internal Models Approach', 'Bank-specific models using Expected Shortfall (MAR31-33)'),
        ('ES', 'Expected Shortfall', 'Average loss in the tail beyond 97.5% VaR; ES_{97.5%} = E[Loss | Loss > VaR_{97.5%}]'),
        ('VaR', 'Value-at-Risk', 'Quantile-based risk measure; replaced by ES in FRTB due to sub-additivity concerns'),
        ('DRC', 'Default Risk Charge', 'Capital for jump-to-default and credit migration risk (MAR22)'),
        ('RRAO', 'Residual Risk Add-On', 'Capital for exotic risks not captured in sensitivity framework (MAR23.5)'),
        ('GIRR', 'General Interest Rate Risk', 'MAR21 risk class covering yield curve, inflation, basis risks'),
        ('CSR', 'Credit Spread Risk', 'MAR21 risk class: CSR_NS (non-securitization), CSR_SEC, CSR_CTP'),
        ('CSR_NS', 'Credit Spread Risk - Non-Securitization', 'Sovereign and corporate bonds/CDS'),
        ('CSR_SEC', 'Credit Spread Risk - Securitization', 'RMBS, CMBS, ABS, CLO tranches'),
        ('CSR_CTP', 'Credit Spread Risk - Correlation Trading Portfolio', 'CDX/iTraxx index tranches'),
        ('NMRF', 'Non-Modellable Risk Factor', 'Risk factor failing modellability test (insufficient data); MAR31.14'),
        ('P&L Attribution', 'Profit and Loss Attribution', 'MAR32 test comparing theoretical vs. actual P&L; eligibility criterion'),
        ('IMCC', 'Internal Model Capital Charge', 'ES capital from IMA; used in 60-day averaging (MAR33.8)'),
        ('RW', 'Risk Weight', 'Prescribed sensitivity shock per MAR21.44-96; e.g., 1.7% for 3m GIRR'),
        ('WS', 'Weighted Sensitivity', 'Net sensitivity × risk weight; WS = s × RW'),
        ('ρ (rho)', 'Intra-bucket correlation', 'Correlation between risk factors within same bucket'),
        ('γ (gamma)', 'Inter-bucket correlation', 'Correlation across different buckets or risk classes'),
        ('Curvature (CVR)', 'Curvature Risk', 'Gamma/convexity risk; incremental loss from non-linearity (MAR21.8)'),
        ('Liquidity Horizon (LH)', 'Liquidity Horizon', 'Time required to exit or hedge position; MAR33 Table 1 (10-120 days)'),
        ('Stressed ES', 'Stressed Expected Shortfall', 'ES calibrated to 12-month stress period; MAR33.6'),
        ('mc', 'IMA Multiplier', 'Backtesting add-on; mc = 1.5 + addon (MAR33.9); ranges 1.5-2.0'),
        ('Backtesting Exception', 'Backtesting Exception', 'Day when actual loss exceeds VaR forecast; triggers mc increase'),
        ('Green Zone', 'Green Zone', '0-4 exceptions in 250 days; mc = 1.5 (no add-on)'),
        ('Amber Zone', 'Amber Zone', '5-9 exceptions; mc = 1.7-1.92 (graduated add-ons)'),
        ('Red Zone', 'Red Zone', '10+ exceptions; mc = 2.0; model revocation risk'),
        ('ICAAP', 'Internal Capital Adequacy Assessment Process', 'Pillar 2 framework; banks\' own capital assessment'),
        ('Pillar 1', 'Pillar 1 (Minimum Capital)', 'Regulatory minimum per standardized formulas'),
        ('Pillar 2', 'Pillar 2 (Supervisory Review)', 'Internal models and stress testing; ICAAP'),
        ('MAR', 'Minimum Capital Requirements for Market Risk', 'Basel III market risk standards (MAR10-99)'),
        ('BCBS', 'Basel Committee on Banking Supervision', 'International regulatory body; sets global banking standards'),
    ]
    
    # Create glossary table
    glossary_table = doc.add_table(rows=1, cols=3)
    glossary_table.style = 'Light Grid Accent 1'
    
    hdr_cells = glossary_table.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Full Name'
    hdr_cells[2].text = 'Definition / Context'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for term, full_name, definition in glossary_data:
        row_cells = glossary_table.add_row().cells
        row_cells[0].text = term
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = full_name
        row_cells[2].text = definition
    
    doc.add_page_break()
    
    # ========================================================================
    # 9. REFERENCES
    # ========================================================================
    
    add_heading(doc, '9. References and Standards', 1)
    
    add_heading(doc, '9.1 Regulatory Documents', 2)
    
    references = [
        'Basel Committee on Banking Supervision (2019). "Minimum capital requirements for market risk." '
        'Bank for International Settlements. (MAR10-99)',
        
        'Basel Committee on Banking Supervision (2016). "Standards: Minimum capital requirements for '
        'market risk." (Original FRTB publication)',
        
        'Basel Committee on Banking Supervision (2013). "Fundamental review of the trading book: '
        'A revised market risk framework." Consultative Document.',
        
        'European Banking Authority (2020). "Guidelines on the specification of the nature, severity '
        'and duration of an economic downturn." (EBA/GL/2020/03)',
        
        'Bank of England (2021). "The Internal Model Approval Process." Prudential Regulation Authority.',
    ]
    
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')
    
    add_heading(doc, '9.2 Academic Literature', 2)
    
    academic_refs = [
        'Christensen, B. J., & Prabhala, N. R. (1998). "The relation between implied and realized '
        'volatility." Journal of Financial Economics, 50(2), 125-150.',
        
        'Forbes, K. J., & Rigobon, R. (2002). "No contagion, only interdependence: measuring stock '
        'market comovements." The Journal of Finance, 57(5), 2223-2261.',
        
        'Gilchrist, S., & Zakrajšek, E. (2012). "Credit spreads and business cycle fluctuations." '
        'American Economic Review, 102(4), 1692-1720.',
        
        'Gordy, M. B., & Howells, B. (2006). "Procyclicality in Basel II: Can we treat the disease '
        'without killing the patient?" Journal of Financial Intermediation, 15(3), 395-417.',
        
        'Whaley, R. E. (2000). "The investor fear gauge." Journal of Portfolio Management, 26(3), 12-17.',
    ]
    
    for ref in academic_refs:
        doc.add_paragraph(ref, style='List Bullet')
    
    add_heading(doc, '9.3 Data Sources', 2)
    
    data_sources = [
        'Federal Reserve Economic Data (FRED). Federal Reserve Bank of St. Louis. https://fred.stlouisfed.org',
        'Bloomberg L.P. Bloomberg Terminal and Server API (BSAPI). https://www.bloomberg.com/professional/support/api-library/',
        'Refinitiv / LSEG. Refinitiv Data Library for Python. https://developers.refinitiv.com/',
        'yfinance Python Library. Yahoo Finance market data wrapper. https://github.com/ranaroussi/yfinance',
        'CBOE Global Markets. VIX Index Methodology. https://www.cboe.com/tradable_products/vix/',
    ]
    
    for source in data_sources:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_page_break()
    
    # ========================================================================
    # 10. DOCUMENT CONTROL
    # ========================================================================
    
    add_heading(doc, '10. Document Control', 1)
    
    add_heading(doc, '10.1 Version History', 2)
    
    version_data = [
        ('1.0', 'April 7, 2026', 'Initial release', 'Comprehensive FRD covering all FRTB components'),
    ]
    
    add_table_from_data(
        doc,
        ['Version', 'Date', 'Author', 'Changes'],
        version_data
    )
    
    doc.add_paragraph()
    
    add_heading(doc, '10.2 Review and Approval', 2)
    
    approval_data = [
        ('Prepared by:', 'Risk Technology', ''),
        ('Reviewed by:', 'Market Risk Management', '[Pending]'),
        ('Approved by:', 'Chief Risk Officer', '[Pending]'),
        ('Next Review Date:', 'October 2026', '(6-month cycle)'),
    ]
    
    add_table_from_data(
        doc,
        ['Role', 'Department', 'Date / Status'],
        approval_data
    )
    
    doc.add_paragraph()
    
    add_heading(doc, '10.3 Distribution', 2)
    
    doc.add_paragraph('This document is distributed to:')
    doc.add_paragraph('• Market Risk Management', style='List Bullet')
    doc.add_paragraph('• Trading Desk Heads', style='List Bullet')
    doc.add_paragraph('• Risk Technology / Quantitative Development', style='List Bullet')
    doc.add_paragraph('• Internal Audit', style='List Bullet')
    doc.add_paragraph('• Model Validation Group', style='List Bullet')
    doc.add_paragraph('• Regulatory Reporting', style='List Bullet')
    
    add_heading(doc, '10.4 Related Documents', 2)
    
    doc.add_paragraph('• FRTB Model Validation Report (separate document)')
    doc.add_paragraph('• P&L Attribution Test Results (MAR32 compliance)')
    doc.add_paragraph('• Backtesting Report (MAR99 quarterly submission)')
    doc.add_paragraph('• FRTB Market Data Enhancement Specification (April 2026)')
    doc.add_paragraph('• ICAAP Stress Testing Framework')
    
    # ========================================================================
    # SAVE DOCUMENT
    # ========================================================================
    
    output_path = '/Users/aaron/Documents/Project/Prometheus/FRTB_Functional_Requirements_Document.docx'
    doc.save(output_path)
    print(f"\n✓ Document generated successfully: {output_path}")
    print(f"  Total pages: ~{len(doc.sections) * 15} (estimated)")
    print(f"  Total sections: 10")
    print(f"  Total tables: {len([elem for elem in doc.element.body if elem.tag.endswith('tbl')])}")
    
    return output_path

if __name__ == "__main__":
    print("="*80)
    print("FRTB Functional Requirements Document Generator")
    print("="*80)
    print("\nGenerating comprehensive DOCX document...")
    print("This covers:")
    print("  1. Regulatory alignment (Basel MAR10-33)")
    print("  2. Business and policy objectives")
    print("  3. Parameter specifications and market risk context")
    print("  4. Process workflows")
    print("  5. Key equations and formulas")
    print("  6. Glossary and references")
    print("\n" + "-"*80)
    
    try:
        output_file = generate_frtb_frd()
        print("\n" + "="*80)
        print("✓ Generation complete!")
        print("="*80)
        print(f"\nOpen with: open {output_file}")
        
    except ImportError as e:
        print(f"\n❌ Error: Missing required library")
        print(f"   {e}")
        print("\nPlease install: pip install python-docx")
    except Exception as e:
        print(f"\n❌ Error during document generation:")
        print(f"   {e}")
        raise
